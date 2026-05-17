"""
API route handlers for Phase 10 Product Runtime.

Routes:
    POST /documents/upload   — upload a document and queue a processing job
    GET  /documents          — list the caller's documents
    GET  /documents/{doc_id} — get a single document (tenant-scoped)
    GET  /jobs               — list the caller's jobs
    GET  /jobs/{job_id}      — get a single job (tenant-scoped)
    POST /queries            — run a knowledge query over the document space
    POST /actions            — execute a legal action workflow (audited)
    GET  /audit              — retrieve the caller's audit trail
"""

from __future__ import annotations

import base64
import uuid
from pathlib import Path
from typing import List

from fastapi import APIRouter, Depends, File, HTTPException, Request, UploadFile, status
from fastapi.responses import FileResponse
from pydantic import BaseModel

from src.actions.action_engine import ActionEngine
from src.actions.action_schema import ActionRequest as EngineActionRequest
from src.api.deps import get_audit, get_runner, get_storage, require_user
from src.api.models import (
    ActionRequest,
    ActionResponse,
    DocumentResponse,
    EvidenceRefOut,
    JobResponse,
    QueryRequest,
    QueryResponse,
    UploadRequest,
    UploadResponse,
)
from src.graphrag.evidence_bundle import EvidenceBundle, empty_bundle
from src.graphrag.reasoning import ReasoningEngine
from src.runtime.audit import AuditLayer
from src.runtime.job_runner import JobRunner
from src.runtime.storage import StorageLayer

router = APIRouter()

_action_engine = ActionEngine()
_reasoning_engine = ReasoningEngine()


# ---------------------------------------------------------------------------
# Auth registration
# ---------------------------------------------------------------------------


class RegisterRequest(BaseModel):
    pass  # user_id is auto-assigned; caller receives the API key


class RegisterResponse(BaseModel):
    user_id: str
    api_key: str


@router.post("/auth/register", response_model=RegisterResponse, status_code=201)
def register(request: Request) -> RegisterResponse:
    """Create a new user and return a fresh API key."""
    storage: StorageLayer = get_storage(request)
    api_key = str(uuid.uuid4())
    record = storage.create_user(api_key)
    return RegisterResponse(user_id=record.user_id, api_key=api_key)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _doc_response(doc) -> DocumentResponse:
    return DocumentResponse(
        doc_id=doc.doc_id,
        filename=doc.filename,
        status=doc.status,
        created_at=doc.created_at,
        metadata=doc.metadata,
    )


def _job_response(job) -> JobResponse:
    return JobResponse(
        job_id=job.job_id,
        doc_id=job.doc_id,
        status=job.status,
        created_at=job.created_at,
        completed_at=job.completed_at,
        error=job.error,
        checkpoint=job.checkpoint,
    )


def _get_bundles(
    request: Request,
    user_id: str,
    document_ids: List[str],
    query: str = "",
    query_id: str = "",
) -> List[EvidenceBundle]:
    """
    Build evidence bundles for a query or action.

    Priority:
    1. app.state.bundle_provider — injected in tests for deterministic results.
    2. app.state.index_store     — real retrieval over stored chunks/graph.
    3. Fallback empty bundle     — when no index is configured.
    """
    provider = getattr(request.app.state, "bundle_provider", None)
    if provider is not None:
        return provider(user_id, document_ids)

    index_store = getattr(request.app.state, "index_store", None)
    if index_store is not None:
        qid = query_id or ("q_" + str(uuid.uuid4())[:8])
        storage: StorageLayer = get_storage(request)
        return index_store.get_bundles(storage, user_id, document_ids, query, qid)

    qid = query_id or ("q_" + str(uuid.uuid4())[:8])
    return [empty_bundle(qid, "Document index not yet built for this document space.")]


# ---------------------------------------------------------------------------
# Upload
# ---------------------------------------------------------------------------


@router.post(
    "/documents/upload",
    response_model=UploadResponse,
    status_code=status.HTTP_202_ACCEPTED,
)
def upload_document(
    body: UploadRequest,
    request: Request,
    user_id: str = Depends(require_user),
) -> UploadResponse:
    storage: StorageLayer = get_storage(request)
    runner: JobRunner = get_runner(request)
    doc = storage.create_document(user_id, body.filename, dict(body.metadata))

    # ── Resolve file location ───────────────────────────────────────────
    if body.file_path:
        # Convenience shortcut: caller supplies a local filesystem path
        resolved = Path(body.file_path)
        if not resolved.exists():
            raise HTTPException(
                status_code=400,
                detail=f"file_path does not exist on the server: {body.file_path}",
            )
        storage.save_file_path(doc.doc_id, str(resolved.resolve()))

    elif body.content_base64:
        # Decode base64 bytes and write to data/uploads/{doc_id}/
        try:
            content = base64.b64decode(body.content_base64)
        except Exception:
            raise HTTPException(status_code=400, detail="content_base64 is not valid base64.")
        upload_dir = Path("data/uploads") / doc.doc_id
        upload_dir.mkdir(parents=True, exist_ok=True)
        dest = upload_dir / body.filename
        dest.write_bytes(content)
        storage.save_file_path(doc.doc_id, str(dest.resolve()))

    # If neither is provided, the job will fail when the processor tries to find the file.
    # This allows metadata-only registration where the file arrives later.

    job = runner.submit(user_id, doc.doc_id)
    return UploadResponse(doc_id=doc.doc_id, job_id=job.job_id)


@router.post(
    "/documents/upload-file",
    response_model=UploadResponse,
    status_code=status.HTTP_202_ACCEPTED,
)
async def upload_document_file(
    request: Request,
    file: UploadFile = File(...),
    user_id: str = Depends(require_user),
) -> UploadResponse:
    """Upload a single document via multipart/form-data (browser-friendly)."""
    storage: StorageLayer = get_storage(request)
    runner: JobRunner = get_runner(request)

    filename = file.filename or "unnamed"
    doc = storage.create_document(user_id, filename)

    upload_dir = Path("data/uploads") / doc.doc_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    dest = upload_dir / filename
    dest.write_bytes(await file.read())
    storage.save_file_path(doc.doc_id, str(dest.resolve()))

    job = runner.submit(user_id, doc.doc_id)
    return UploadResponse(doc_id=doc.doc_id, job_id=job.job_id)


# ---------------------------------------------------------------------------
# Documents
# ---------------------------------------------------------------------------


@router.get("/documents", response_model=List[DocumentResponse])
def list_documents(
    request: Request,
    user_id: str = Depends(require_user),
) -> List[DocumentResponse]:
    storage: StorageLayer = get_storage(request)
    return [_doc_response(d) for d in storage.list_documents(user_id)]


@router.get("/documents/{doc_id}", response_model=DocumentResponse)
def get_document(
    doc_id: str,
    request: Request,
    user_id: str = Depends(require_user),
) -> DocumentResponse:
    storage: StorageLayer = get_storage(request)
    doc = storage.get_document(user_id, doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found.")
    return _doc_response(doc)


# ---------------------------------------------------------------------------
# Contract text extraction (synchronous — for frontend contract analysis)
# ---------------------------------------------------------------------------


@router.post("/contract/extract-text")
async def extract_contract_text(
    file: UploadFile = File(...),
) -> dict:
    """
    Extract plain text from an uploaded PDF, DOCX, DOC, or TXT file.
    Returns {text, filename} synchronously for use in the contract analysis UI.
    """
    import os
    import tempfile

    filename = file.filename or "upload"
    suffix = Path(filename).suffix.lower()
    content = await file.read()

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    extracted = ""
    error_msg = ""
    try:
        if suffix == ".txt":
            extracted = content.decode("utf-8", errors="replace")
        elif suffix == ".docx":
            try:
                import docx as python_docx  # python-docx
                doc = python_docx.Document(tmp_path)
                extracted = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except Exception as e:
                error_msg = f"Không thể đọc DOCX: {e}"
        elif suffix == ".pdf":
            try:
                from pdfminer.high_level import extract_text as _pdf_extract
                extracted = _pdf_extract(tmp_path) or ""
            except Exception as e:
                error_msg = f"Không thể đọc PDF: {e}"
        elif suffix == ".doc":
            try:
                import docx2txt
                extracted = docx2txt.process(tmp_path) or ""
            except Exception as e:
                error_msg = f"Không thể đọc DOC: {e}"
        else:
            error_msg = f"Định dạng '{suffix}' không được hỗ trợ. Vui lòng dùng TXT, DOCX hoặc PDF."
    finally:
        os.unlink(tmp_path)

    if error_msg:
        return {"text": "", "filename": filename, "error": error_msg}
    return {"text": extracted[:60000], "filename": filename}


# ---------------------------------------------------------------------------
# Jobs
# ---------------------------------------------------------------------------


@router.get("/jobs", response_model=List[JobResponse])
def list_jobs(
    request: Request,
    user_id: str = Depends(require_user),
) -> List[JobResponse]:
    storage: StorageLayer = get_storage(request)
    return [_job_response(j) for j in storage.list_jobs(user_id)]


@router.get("/jobs/{job_id}", response_model=JobResponse)
def get_job(
    job_id: str,
    request: Request,
    user_id: str = Depends(require_user),
) -> JobResponse:
    storage: StorageLayer = get_storage(request)
    job = storage.get_job(user_id, job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found.")
    return _job_response(job)


# ---------------------------------------------------------------------------
# Queries
# ---------------------------------------------------------------------------


@router.post("/queries", response_model=QueryResponse)
def run_query(
    body: QueryRequest,
    request: Request,
    user_id: str = Depends(require_user),
) -> QueryResponse:
    query_id = body.query_id or ("q_" + str(uuid.uuid4())[:8])
    bundles = _get_bundles(request, user_id, body.document_ids, query=body.query, query_id=query_id)
    primary = bundles[0] if bundles else empty_bundle(query_id, "No documents.")
    result = _reasoning_engine.reason(body.query, primary, query_id=query_id)

    return QueryResponse(
        query_id=result.query_id,
        query=result.query,
        intent=result.intent,
        support_status=result.support_status,
        confidence=result.confidence,
        answer=result.answer,
        citations=result.citations,
        evidence_excerpts=result.evidence_excerpts,
        is_refusal=result.is_refusal,
        warnings=result.warnings,
    )


# ---------------------------------------------------------------------------
# Actions
# ---------------------------------------------------------------------------


@router.post("/actions", response_model=ActionResponse)
def run_action(
    body: ActionRequest,
    request: Request,
    user_id: str = Depends(require_user),
    audit: AuditLayer = Depends(get_audit),
) -> ActionResponse:
    request_id = body.request_id or ("req_" + str(uuid.uuid4())[:8])

    try:
        engine_req = EngineActionRequest(
            action_type=body.action_type,
            query=body.query,
            document_ids=body.document_ids,
            parameters=body.parameters,
            request_id=request_id,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))

    bundles = _get_bundles(request, user_id, body.document_ids, query=body.query, query_id=request_id)
    result = _action_engine.execute(engine_req, bundles)
    audit_record = audit.log_action_result(user_id, result)

    return ActionResponse(
        request_id=result.request_id,
        action_type=result.action_type,
        status=result.status,
        output=result.output,
        citations=result.citations,
        evidence_refs=[
            EvidenceRefOut(
                chunk_id=r.chunk_id,
                citation=r.citation,
                excerpt=r.excerpt,
                confidence=r.confidence,
                is_generated=r.is_generated,
            )
            for r in result.evidence_refs
        ],
        assumptions=result.assumptions,
        missing_evidence=result.missing_evidence,
        warnings=result.warnings,
        is_generated=result.is_generated,
        refusal_reason=result.refusal_reason,
        audit_id=audit_record.audit_id,
    )


# ---------------------------------------------------------------------------
# Audit trail
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Document content viewer & download (Phase 14)
# ---------------------------------------------------------------------------


def _get_vs(request: Request):
    """Return VectorStorage from app.state, or None if MongoDB not available."""
    return getattr(request.app.state, "vector_storage", None)


@router.get("/documents/{doc_id}/content")
def get_document_content(
    doc_id: str,
    request: Request,
    user_id: str = Depends(require_user),
) -> dict:
    """Return extracted text (chunks) for an uploaded document."""
    storage: StorageLayer = get_storage(request)
    doc = storage.get_document(user_id, doc_id)
    if doc is None:
        # fall back to global doc lookup
        doc = storage.get_document_by_id(doc_id)
        if doc is None or not doc.is_global:
            raise HTTPException(status_code=404, detail="Document not found.")

    vs = _get_vs(request)
    chunks: list = []
    if vs is not None:
        chunks = vs.get_chunks_by_document(doc_id, user_id)

    extracted_text = "\n\n---\n\n".join(c.get("content", "") for c in chunks)
    law_type = chunks[0].get("law_type", "") if chunks else ""

    return {
        "doc_id": doc_id,
        "filename": doc.filename,
        "status": doc.status,
        "chunk_count": len(chunks),
        "law_type": law_type,
        "extracted_text": extracted_text,
    }


@router.get("/documents/{doc_id}/download")
def download_document(
    doc_id: str,
    request: Request,
    user_id: str = Depends(require_user),
) -> FileResponse:
    """Serve the original uploaded file for download."""
    storage: StorageLayer = get_storage(request)
    doc = storage.get_document(user_id, doc_id)
    if doc is None:
        doc = storage.get_document_by_id(doc_id)
        if doc is None or not doc.is_global:
            raise HTTPException(status_code=404, detail="Document not found.")

    file_path = storage.get_file_path(doc_id)
    if file_path is None or not Path(file_path).exists():
        raise HTTPException(status_code=404, detail="File not found on disk.")

    return FileResponse(
        path=file_path,
        filename=doc.filename,
        media_type="application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{doc.filename}"'},
    )


# ---------------------------------------------------------------------------
# Session evidence upload (Phase 14)
# ---------------------------------------------------------------------------


@router.post("/sessions/{session_id}/evidence")
async def upload_session_evidence(
    session_id: str,
    request: Request,
    file: UploadFile = File(...),
    user_id: str = Depends(require_user),
) -> dict:
    """
    Upload additional evidence (PDF/DOCX/TXT) into an existing analysis session.
    Extracted text is appended to the session context so the next /intelligence/analyze
    call sees it without the user having to paste content manually.
    """
    import os
    import tempfile
    import uuid as _uuid

    filename = file.filename or "evidence"
    suffix = Path(filename).suffix.lower()
    content = await file.read()

    if suffix not in {".pdf", ".docx", ".doc", ".txt"}:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{suffix}'. Use PDF, DOCX, DOC, or TXT.",
        )

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    extracted = ""
    try:
        if suffix == ".txt":
            extracted = content.decode("utf-8", errors="replace")
        elif suffix == ".docx":
            import docx as python_docx
            doc = python_docx.Document(tmp_path)
            extracted = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif suffix == ".pdf":
            from pdfminer.high_level import extract_text as _pdf_extract
            extracted = _pdf_extract(tmp_path) or ""
        elif suffix == ".doc":
            import docx2txt
            extracted = docx2txt.process(tmp_path) or ""
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Could not extract text: {exc}")
    finally:
        os.unlink(tmp_path)

    snippet = extracted[:2000].strip()
    evidence_id = "ev_" + str(_uuid.uuid4())[:8]

    # Persist into MongoDB session so orchestrator picks it up on next query
    try:
        from src.memory.session_store import SessionStore
        SessionStore().append_evidence(session_id, {
            "evidence_id": evidence_id,
            "filename": filename,
            "snippet": snippet,
            "char_count": len(extracted),
        })
    except Exception:
        pass  # best-effort; don't fail the upload if session store is unavailable

    return {
        "evidence_id": evidence_id,
        "session_id": session_id,
        "filename": filename,
        "snippet": snippet[:300],
        "char_count": len(extracted),
        "status": "attached",
    }


@router.get("/audit")
def get_audit_trail(
    request: Request,
    user_id: str = Depends(require_user),
    audit: AuditLayer = Depends(get_audit),
):
    records = audit.get_trail(user_id)
    return [
        {
            "audit_id": r.audit_id,
            "request_id": r.request_id,
            "action_type": r.action_type,
            "status": r.status,
            "output_hash": r.output_hash,
            "created_at": r.created_at,
        }
        for r in records
    ]
