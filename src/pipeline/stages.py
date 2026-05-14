"""
Pipeline stage implementations.

Each function implements one stage of the pipeline and returns a StageOutput.
Stages are deterministic, local-first, and hallucination-safe by default.

Implementation philosophy:
- This first version uses local Python libraries (pdfminer, python-docx, BeautifulSoup).
- AI-assisted repair is gated by config.enable_ai_repair and never runs by default.
- Each stage is a plain function: easy to test, replace, or extend.
- When a library is not installed, the stage degrades gracefully and emits a warning.

Language support:
- Vietnamese (vi) and English (en) legal hierarchies are both recognised.
- Vietnamese hierarchy: Chương → Mục → Điều → Khoản → Điểm
- English hierarchy:    Chapter → Section → Article → Clause → Point
- All text is handled as UTF-8. Vietnamese diacritics are preserved exactly.
- Language detection is heuristic; override via metadata hints if known.

File format support:
- .pdf  — pdfminer (local text layer) or OCR fallback
- .docx — python-docx
- .doc  — docx2txt (plain text) → python-docx (if .doc can be opened) → antiword CLI
- .html — beautifulsoup4
- image — pytesseract OCR
"""

from __future__ import annotations

import hashlib
import mimetypes
import re
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from src.config import PipelineConfig
from src.pipeline.interfaces import StageContext, StageOutput
from src.schemas.chunk import Chunk, ChunkSet, ChunkingDecision
from src.schemas.document import (
    Article,
    Block,
    CanonicalDocument,
    Clause,
    Confidence,
    DocumentMetadata,
    DocumentProfile,
    Image,
    LanguageDetection,
    RenderedOutputs,
    Section,
    Table,
    TableCell,
    Traceability,
    ValidationSummary,
)
from src.schemas.evaluation import STATUS_FAIL, STATUS_PASS, STATUS_WARNING
from src.schemas.graph import GraphEdge, GraphNode, GraphSubgraph, GraphBuildJob
from src.utils import trace as T

# Multilingual retrieval modules
from src.retrieval.language_detector import detect_language as _formal_detect_language
from src.retrieval.canonical_references import CanonicalRefBuilder, canonical_refs_for_chunk
from src.retrieval.retrieval_engine import RetrievalEngine
from src.graphrag.legal_ontology import add_alias_edges


# ---------------------------------------------------------------------------
# Stage 1: Input loading
# ---------------------------------------------------------------------------


def stage_input_loading(ctx: StageContext) -> StageOutput:
    """
    Validate the source file, detect file type, collect basic metadata.
    Computes source_hash for provenance.
    """
    path = ctx.source_path
    warnings: List[str] = []
    errors: List[str] = []

    if not path.exists():
        return StageOutput(
            stage_name="input_loading",
            status=STATUS_FAIL,
            summary=f"Source file not found: {path}",
            errors=[f"File not found: {path}"],
        )

    # Compute source hash
    source_hash = T.compute_source_hash(path)

    # Detect MIME type and file extension
    mime, _ = mimetypes.guess_type(str(path))
    suffix = path.suffix.lower().lstrip(".")

    file_type_map = {
        "pdf": "pdf",
        "docx": "docx",
        "doc": "doc",       # legacy Word — handled by its own extractor
        "html": "html",
        "htm": "html",
        "png": "image",
        "jpg": "image",
        "jpeg": "image",
        "tiff": "image",
        "tif": "image",
        "bmp": "image",
        "webp": "image",
    }
    file_type = file_type_map.get(suffix, "unknown")

    if file_type == "unknown":
        warnings.append(f"Unknown file type for extension '.{suffix}'; will attempt text extraction.")

    file_size_bytes = path.stat().st_size
    ctx.put("source_hash", source_hash)
    ctx.put("file_type", file_type)
    ctx.put("mime_type", mime or "application/octet-stream")
    ctx.put("file_size_bytes", file_size_bytes)

    ctx.logger.info(
        f"Loaded: {path.name} | type={file_type} | size={file_size_bytes:,}B | hash={source_hash[:12]}",
        stage="input_loading",
    )

    return StageOutput(
        stage_name="input_loading",
        status=STATUS_WARNING if warnings else STATUS_PASS,
        summary=f"Loaded '{path.name}' (type={file_type}, size={file_size_bytes:,} bytes)",
        warnings=warnings,
        errors=errors,
        output_summary={
            "filename": path.name,
            "file_type": file_type,
            "file_size_bytes": file_size_bytes,
            "source_hash": source_hash,
        },
    )


# ---------------------------------------------------------------------------
# Stage 2: Document profiling
# ---------------------------------------------------------------------------


def stage_document_profiling(ctx: StageContext) -> StageOutput:
    """
    Profile the document to determine extraction strategy.
    Local-first: uses file introspection without AI.
    """
    path = ctx.source_path
    file_type = ctx.get("file_type", "unknown")
    cfg = ctx.config
    warnings: List[str] = []

    profile = DocumentProfile(file_type=file_type)

    try:
        if file_type == "pdf":
            profile = _profile_pdf(path, cfg, warnings)
        elif file_type == "docx":
            profile = _profile_docx(path, cfg, warnings)
        elif file_type == "doc":
            profile = _profile_doc(path, cfg, warnings)
        elif file_type == "html":
            profile = _profile_html(path, cfg, warnings)
        elif file_type == "image":
            profile = _profile_image(path, cfg, warnings)
        else:
            warnings.append(f"Cannot profile unknown file type '{file_type}'; using defaults.")
            profile.extraction_strategy = "simple_local"
    except Exception as exc:
        warnings.append(f"Profiling error: {exc}. Using conservative defaults.")
        profile.extraction_strategy = "simple_local"

    ctx.put("profile", profile)
    ctx.logger.decision(
        stage="document_profiling",
        decision_type="extraction_strategy",
        value=profile.extraction_strategy,
        reason=(
            f"text_layer={profile.text_layer_coverage:.2f}, "
            f"scan_quality={profile.scan_quality_score:.2f}, "
            f"layout_complexity={profile.layout_complexity_score:.2f}, "
            f"pages={profile.page_count}, "
            f"long_doc={profile.is_long_document}"
        ),
        thresholds={
            "strong_text_layer": cfg.strong_text_layer_threshold,
            "weak_text_layer": cfg.weak_text_layer_threshold,
            "long_document_pages": cfg.long_document_page_threshold,
        },
    )

    return StageOutput(
        stage_name="document_profiling",
        status=STATUS_WARNING if warnings else STATUS_PASS,
        summary=f"Strategy: {profile.extraction_strategy} | pages={profile.page_count} | tables={profile.has_tables}",
        warnings=warnings,
        output_summary={
            "extraction_strategy": profile.extraction_strategy,
            "page_count": profile.page_count,
            "text_layer_coverage": round(profile.text_layer_coverage, 3),
            "is_long_document": profile.is_long_document,
            "has_tables": profile.has_tables,
            "has_images": profile.has_images,
        },
    )


def _profile_pdf(path: Path, cfg: PipelineConfig, warnings: List[str]) -> DocumentProfile:
    """Profile a PDF using pdfminer or pypdf if available."""
    profile = DocumentProfile(file_type="pdf")
    try:
        import pdfminer.high_level as pdfminer
        from pdfminer.layout import LAParams

        text = pdfminer.extract_text(str(path), laparams=LAParams())
        page_count = max(1, text.count("\x0c") + 1)
        text_length = len(text.strip())

        profile.page_count = page_count
        profile.text_layer_coverage = min(1.0, text_length / max(1, page_count * 500))
        profile.estimated_token_count = text_length // 4  # rough token estimate

        # Heuristic table detection
        profile.has_tables = bool(re.search(r"\|.*\||\t.*\t", text))
        profile.is_long_document = (
            page_count >= cfg.long_document_page_threshold
            or profile.estimated_token_count >= cfg.long_document_token_threshold
        )
        profile.is_scan_heavy = profile.text_layer_coverage < cfg.weak_text_layer_threshold
        # Language detection (heuristic + formal)
        profile.languages = _detect_languages(text)
        _attach_language_detection(profile, text)

    except ImportError:
        warnings.append("pdfminer not installed; using minimal PDF profiling.")
        profile.page_count = 1
        profile.text_layer_coverage = 0.5
        profile.extraction_strategy = "simple_local"
        return profile

    # Route extraction strategy
    profile.extraction_strategy = _choose_strategy(profile, cfg)
    return profile


def _profile_docx(path: Path, cfg: PipelineConfig, warnings: List[str]) -> DocumentProfile:
    profile = DocumentProfile(file_type="docx")
    try:
        import docx  # python-docx

        doc = docx.Document(str(path))
        para_text = " ".join(p.text for p in doc.paragraphs)
        profile.page_count = max(1, len(para_text) // 3000)
        profile.text_layer_coverage = 1.0
        profile.has_tables = len(doc.tables) > 0
        profile.estimated_token_count = len(para_text) // 4
        profile.is_long_document = profile.estimated_token_count >= cfg.long_document_token_threshold
        profile.languages = _detect_languages(para_text)
        _attach_language_detection(profile, para_text)
    except ImportError:
        warnings.append("python-docx not installed; using fallback DOCX profiling.")
        profile.text_layer_coverage = 1.0

    profile.extraction_strategy = "simple_local" if not profile.is_long_document else "long_local"
    return profile


def _profile_doc(path: Path, cfg: PipelineConfig, warnings: List[str]) -> DocumentProfile:
    """
    Profile a legacy .doc file.
    Tries docx2txt first (pure Python), then falls back to antiword CLI.
    """
    profile = DocumentProfile(file_type="doc")
    text = _read_doc_text(path, warnings)
    if text:
        profile.text_layer_coverage = 1.0
        profile.estimated_token_count = len(text) // 4
        profile.has_tables = False  # docx2txt loses table topology; mark as missing
        profile.is_long_document = profile.estimated_token_count >= cfg.long_document_token_threshold
        # Detect Vietnamese or English
        profile.languages = _detect_languages(text)
        _attach_language_detection(profile, text)
    else:
        warnings.append(".doc file produced no text. File may be corrupt or binary-only.")
        profile.text_layer_coverage = 0.0
    profile.extraction_strategy = "simple_local" if not profile.is_long_document else "long_local"
    return profile


def _read_doc_text(path: Path, warnings: List[str]) -> str:
    """
    Extract plain text from a legacy .doc file.
    Priority: docx2txt → antiword CLI → empty string with warning.
    """
    # Attempt 1: docx2txt (works on both .doc and .docx in many cases)
    try:
        import docx2txt
        text = docx2txt.process(str(path))
        if text and text.strip():
            return text
    except ImportError:
        pass
    except Exception:
        pass

    # Attempt 2: antiword (command-line tool, must be installed on the OS)
    try:
        import subprocess
        result = subprocess.run(
            ["antiword", str(path)],
            capture_output=True, text=True, timeout=30,
            encoding="utf-8", errors="replace",
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except FileNotFoundError:
        pass
    except Exception:
        pass

    # Attempt 3: python-docx (occasionally opens .doc if saved as OOXML)
    try:
        import docx
        doc = docx.Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if text.strip():
            return text
    except Exception:
        pass

    warnings.append(
        f"Could not extract text from '{path.name}'. "
        "Install docx2txt (`pip install docx2txt`) or antiword (OS package) for .doc support."
    )
    return ""


def _profile_html(path: Path, cfg: PipelineConfig, warnings: List[str]) -> DocumentProfile:
    profile = DocumentProfile(file_type="html")
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        profile.text_layer_coverage = 1.0
        # HTML files are single logical "page" — use token count for page estimate
        char_len = len(text)
        profile.estimated_token_count = char_len // 4
        profile.page_count = max(1, profile.estimated_token_count // 2000)  # ~2000 tokens/page
        profile.has_tables = "<table" in text.lower()
        profile.has_images = "<img" in text.lower()
        profile.is_long_document = profile.estimated_token_count >= cfg.long_document_token_threshold
        profile.languages = _detect_languages(text)
        _attach_language_detection(profile, text)
    except Exception as exc:
        warnings.append(f"HTML read error: {exc}")
        profile.page_count = 1
    profile.extraction_strategy = "simple_local"
    return profile


def _profile_image(path: Path, cfg: PipelineConfig, warnings: List[str]) -> DocumentProfile:
    profile = DocumentProfile(file_type="image")
    profile.text_layer_coverage = 0.0
    profile.is_scan_heavy = True
    profile.image_density = 1.0
    if not cfg.enable_ocr:
        warnings.append("Image file detected but OCR is disabled; text extraction will be empty.")
    profile.extraction_strategy = "scan_recovery"
    profile.page_count = 1
    return profile


def _attach_language_detection(profile: DocumentProfile, text: str) -> None:
    """
    Run the formal language detector and attach the result to the document profile.
    Uses the first 6000 chars for performance; updates profile.language_detection in-place.
    Also syncs profile.languages from the detection result when confident (>= 0.6).
    """
    result = _formal_detect_language(text[:6000])
    profile.language_detection = LanguageDetection(
        language=result.language,
        confidence=result.confidence,
        jurisdiction=result.jurisdiction,
        script=result.script,
        signals=result.signals[:15],
    )
    # Sync profile.languages if detection is confident
    if result.confidence >= 0.6 and result.language not in ("unknown", "mixed"):
        if result.language not in profile.languages:
            profile.languages = [result.language] + [
                l for l in profile.languages if l != result.language
            ]
    elif result.language == "mixed":
        # Ensure both vi and en are present for mixed documents
        for lang in ("vi", "en"):
            if lang not in profile.languages:
                profile.languages.append(lang)


def _detect_languages(text: str) -> List[str]:
    """
    Heuristic language detection for Vietnamese and English.
    Vietnamese is detected by the presence of diacritical characters unique to it.
    Returns a list of detected language codes (e.g. ["vi"], ["en"], ["vi", "en"]).
    """
    # Vietnamese-specific diacritics (tonal marks and vowel modifiers)
    vi_pattern = re.compile(
        r"[àáâãèéêìíòóôõùúýăđơưạảấầẩẫậắằẳẵặẹẻẽếềểễệỉịọỏốồổỗộớờởỡợụủứừửữựỳỷỹỵ"
        r"ÀÁÂÃÈÉÊÌÍÒÓÔÕÙÚÝĂĐƠƯẠẢẤẦẨẪẬẮẰẲẴẶẸẺẼẾỀỂỄỆỈỊỌỎỐỒỔỖỘỚỜỞỠỢỤỦỨỪỬỮỰỲỶỸỴ]",
        re.UNICODE,
    )
    # Vietnamese legal keywords
    vi_keywords = re.compile(
        r"\b(điều|khoản|điểm|chương|mục|nghị định|thông tư|luật|bộ luật|quyết định|nghị quyết)\b",
        re.IGNORECASE | re.UNICODE,
    )
    # English legal keywords
    en_keywords = re.compile(
        r"\b(article|clause|section|chapter|whereas|herein|agreement|contract|whereas|shall|obligation)\b",
        re.IGNORECASE,
    )

    langs = []
    sample = text[:5000]  # check first 5000 chars for speed

    has_vi_chars = bool(vi_pattern.search(sample))
    has_vi_keywords = bool(vi_keywords.search(sample))
    has_en_keywords = bool(en_keywords.search(sample))

    if has_vi_chars or has_vi_keywords:
        langs.append("vi")
    if has_en_keywords:
        langs.append("en")
    if not langs:
        langs.append("en")  # safe default
    return langs


def _choose_strategy(profile: DocumentProfile, cfg: PipelineConfig) -> str:
    """
    Routing logic from docs/parsing/document-intelligence-pipeline.md decision tree.
    """
    if profile.file_type in ("docx", "html"):
        return "simple_local" if not profile.is_long_document else "long_local"
    if profile.text_layer_coverage >= cfg.strong_text_layer_threshold:
        if profile.is_long_document:
            return "long_local"
        return "simple_local"
    if profile.text_layer_coverage < cfg.weak_text_layer_threshold:
        return "scan_recovery"
    return "hybrid_region_precision"


# ---------------------------------------------------------------------------
# Stage 3: Extraction
# ---------------------------------------------------------------------------


def stage_extraction(ctx: StageContext) -> StageOutput:
    """
    Extract text, tables, and images from the source document.
    Local-first; AI repair only if explicitly enabled.
    """
    path = ctx.source_path
    profile: DocumentProfile = ctx.get("profile", DocumentProfile())
    cfg = ctx.config
    warnings: List[str] = []
    errors: List[str] = []

    raw_blocks: List[Dict[str, Any]] = []
    raw_tables: List[Dict[str, Any]] = []
    raw_images: List[Dict[str, Any]] = []
    page_count = profile.page_count or 1

    try:
        if profile.file_type == "pdf":
            raw_blocks, raw_tables = _extract_pdf(path, profile, cfg, warnings)
        elif profile.file_type == "docx":
            raw_blocks, raw_tables = _extract_docx(path, profile, cfg, warnings)
        elif profile.file_type == "doc":
            raw_blocks, raw_tables = _extract_doc(path, profile, cfg, warnings)
        elif profile.file_type == "html":
            raw_blocks, raw_tables = _extract_html(path, profile, cfg, warnings)
        elif profile.file_type == "image":
            raw_blocks = _extract_image(path, profile, cfg, warnings)
        else:
            warnings.append(f"No extractor for file_type='{profile.file_type}'. Returning empty blocks.")
    except Exception as exc:
        errors.append(f"Extraction failed: {exc}")
        ctx.logger.error("Extraction error", stage="extraction", error=str(exc))

    ctx.put("raw_blocks", raw_blocks)
    ctx.put("raw_tables", raw_tables)
    ctx.put("raw_images", raw_images)

    return StageOutput(
        stage_name="extraction",
        status=STATUS_FAIL if errors else (STATUS_WARNING if warnings else STATUS_PASS),
        summary=f"Extracted {len(raw_blocks)} blocks, {len(raw_tables)} tables from {page_count} pages",
        warnings=warnings,
        errors=errors,
        output_summary={
            "raw_block_count": len(raw_blocks),
            "raw_table_count": len(raw_tables),
            "raw_image_count": len(raw_images),
        },
    )


def _extract_pdf(
    path: Path, profile: DocumentProfile, cfg: PipelineConfig, warnings: List[str]
) -> Tuple[List[Dict], List[Dict]]:
    blocks: List[Dict] = []
    tables: List[Dict] = []
    try:
        import pdfminer.high_level as pdfminer
        from pdfminer.layout import LAParams, LTTextBox, LTFigure

        for page_num, page_layout in enumerate(
            pdfminer.extract_pages(str(path), laparams=LAParams()), start=1
        ):
            for element_idx, element in enumerate(page_layout):
                if hasattr(element, "get_text"):
                    text = element.get_text().strip()
                    if text:
                        blocks.append({
                            "page_index": page_num,
                            "element_index": element_idx,
                            "raw_text": text,
                            "bbox": getattr(element, "bbox", None),
                            "source": "pdfminer",
                        })
    except ImportError:
        warnings.append("pdfminer not installed; PDF text extraction skipped.")
    except Exception as exc:
        warnings.append(f"PDF extraction error: {exc}")
    return blocks, tables


def _extract_docx(
    path: Path, profile: DocumentProfile, cfg: PipelineConfig, warnings: List[str]
) -> Tuple[List[Dict], List[Dict]]:
    blocks: List[Dict] = []
    tables: List[Dict] = []
    try:
        import docx

        doc = docx.Document(str(path))
        for para_idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                blocks.append({
                    "page_index": 1,
                    "element_index": para_idx,
                    "raw_text": para.text.strip(),
                    "style": para.style.name if para.style else "",
                    "source": "python-docx",
                })
        for tbl_idx, tbl in enumerate(doc.tables):
            rows = []
            for row in tbl.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append({
                "page_index": 1,
                "table_index": tbl_idx,
                "rows": rows,
                "source": "python-docx",
            })
    except ImportError:
        warnings.append("python-docx not installed; DOCX extraction skipped.")
    except Exception as exc:
        warnings.append(f"DOCX extraction error: {exc}")
    return blocks, tables


def _extract_doc(
    path: Path, profile: DocumentProfile, cfg: PipelineConfig, warnings: List[str]
) -> Tuple[List[Dict], List[Dict]]:
    """
    Extract text from a legacy .doc file.
    Uses the same _read_doc_text() helper as profiling.
    Tables: docx2txt does not preserve table topology, so we emit the
    raw text as blocks and add a warning about table loss.
    """
    blocks: List[Dict] = []
    tables: List[Dict] = []

    text = _read_doc_text(path, warnings)
    if not text:
        return blocks, tables

    # Split into paragraphs on blank lines or line breaks
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    for idx, para in enumerate(paragraphs):
        blocks.append({
            "page_index": 1,
            "element_index": idx,
            "raw_text": para,
            "source": "doc_text_extraction",
        })

    warnings.append(
        f"'.doc' format: table topology cannot be reliably preserved with text extraction. "
        "Convert to .docx for full table support."
    )
    return blocks, tables


def _extract_html(
    path: Path, profile: DocumentProfile, cfg: PipelineConfig, warnings: List[str]
) -> Tuple[List[Dict], List[Dict]]:
    blocks: List[Dict] = []
    tables: List[Dict] = []
    try:
        from bs4 import BeautifulSoup

        html_text = path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(html_text, "html.parser")

        # Extract paragraphs and headings
        for el_idx, tag in enumerate(soup.find_all(["p", "h1", "h2", "h3", "h4", "h5", "li"])):
            text = tag.get_text(strip=True)
            if text:
                blocks.append({
                    "page_index": 1,
                    "element_index": el_idx,
                    "raw_text": text,
                    "tag": tag.name,
                    "source": "beautifulsoup",
                })

        # Extract tables
        for tbl_idx, tbl in enumerate(soup.find_all("table")):
            rows = []
            for row in tbl.find_all("tr"):
                rows.append([cell.get_text(strip=True) for cell in row.find_all(["td", "th"])])
            tables.append({
                "page_index": 1,
                "table_index": tbl_idx,
                "rows": rows,
                "source": "beautifulsoup",
            })
    except ImportError:
        # Fallback: plain text extraction
        warnings.append("beautifulsoup4 not installed; falling back to plain HTML text extraction.")
        text = path.read_text(encoding="utf-8", errors="replace")
        # Strip tags crudely
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        if text:
            blocks.append({
                "page_index": 1,
                "element_index": 0,
                "raw_text": text,
                "source": "regex_strip",
            })
    except Exception as exc:
        warnings.append(f"HTML extraction error: {exc}")
    return blocks, tables


def _extract_image(
    path: Path, profile: DocumentProfile, cfg: PipelineConfig, warnings: List[str]
) -> List[Dict]:
    blocks: List[Dict] = []
    if not cfg.enable_ocr:
        warnings.append("OCR disabled; image text extraction skipped.")
        return blocks
    try:
        import pytesseract
        from PIL import Image as PILImage

        img = PILImage.open(str(path))
        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        words = [
            w for w, conf in zip(data["text"], data["conf"])
            if w.strip() and int(conf) > 0
        ]
        confidences = [int(c) for c in data["conf"] if int(c) > 0]
        avg_conf = sum(confidences) / len(confidences) if confidences else 0
        text = " ".join(words)
        if text:
            blocks.append({
                "page_index": 1,
                "element_index": 0,
                "raw_text": text,
                "ocr_confidence": avg_conf / 100.0,
                "source": "tesseract",
            })
    except ImportError:
        warnings.append("pytesseract or Pillow not installed; OCR skipped for image input.")
    except Exception as exc:
        warnings.append(f"OCR error: {exc}")
    return blocks


# ---------------------------------------------------------------------------
# Stage 4: Canonical structuring
# ---------------------------------------------------------------------------


def stage_canonical_structuring(ctx: StageContext) -> StageOutput:
    """
    Convert raw extraction output into a CanonicalDocument with provenance.
    Detects headings, articles, clauses. Preserves tables and images as first-class objects.
    """
    raw_blocks: List[Dict] = ctx.get("raw_blocks", [])
    raw_tables: List[Dict] = ctx.get("raw_tables", [])
    raw_images: List[Dict] = ctx.get("raw_images", [])
    profile: DocumentProfile = ctx.get("profile", DocumentProfile())
    source_hash = ctx.get("source_hash", "")
    file_type = ctx.get("file_type", "unknown")
    cfg = ctx.config

    # Build traceability template
    traceability_base = Traceability(
        trace_id=ctx.trace_id,
        source_hash=source_hash,
        extractor="local_pipeline_v1",
        extractor_version=cfg.processing_version,
        processing_version=cfg.processing_version,
        created_at=T.now_iso(),
    )

    # Detect document language from profile (set during profiling)
    profile_langs = profile.languages or ["en"]
    primary_language = profile_langs[0] if profile_langs else "en"

    # Build Block objects
    blocks: List[Block] = []
    for idx, rb in enumerate(raw_blocks):
        page_index = rb.get("page_index", 1)
        page_id = T.make_page_id(ctx.document_id, page_index)
        region_id = T.make_region_id(page_id, idx, "text")
        block_id = T.make_block_id(ctx.document_id, page_index, 0, idx)
        raw_text = rb.get("raw_text", "")
        block_type = _detect_block_type(rb)
        ocr_conf = rb.get("ocr_confidence")
        overall_conf = ocr_conf if ocr_conf is not None else 1.0
        degraded = overall_conf < cfg.extraction_confidence_threshold

        block = Block(
            block_id=block_id,
            page_id=page_id,
            region_id=region_id,
            block_type=block_type,
            order_index=idx,
            raw_text=raw_text,
            clean_text=_clean_text(raw_text),
            language=primary_language,
            confidence=Confidence(
                overall=overall_conf,
                extraction=overall_conf,
                ocr=ocr_conf,
                degraded=degraded,
                reasons=["ocr_below_threshold"] if degraded and ocr_conf is not None else [],
            ),
            traceability=traceability_base,
        )
        blocks.append(block)

    # Build Table objects
    tables: List[Table] = []
    for tbl_idx, rt in enumerate(raw_tables):
        page_index = rt.get("page_index", 1)
        table_id = T.make_table_id(ctx.document_id, page_index, tbl_idx)
        rows_raw = rt.get("rows", [])
        page_id = T.make_page_id(ctx.document_id, page_index)

        cells: List[TableCell] = []
        for r_idx, row in enumerate(rows_raw):
            for c_idx, cell_text in enumerate(row):
                cells.append(TableCell(row=r_idx, col=c_idx, text=cell_text, raw_text=cell_text))

        # Simple markdown projection for retrieval
        md_lines = []
        for row in rows_raw:
            md_lines.append("| " + " | ".join(str(c) for c in row) + " |")
        projection = "\n".join(md_lines)

        tables.append(Table(
            table_id=table_id,
            page_refs=[page_id],
            row_count=len(rows_raw),
            column_count=max((len(r) for r in rows_raw), default=0),
            cells=cells,
            projection_text=projection,
            markdown=projection,
            confidence=Confidence(overall=0.9, topology=0.9),
            traceability=traceability_base,
        ))

    # Build Image objects (from raw_images if any, else empty)
    images: List[Image] = []
    for img_idx, ri in enumerate(raw_images):
        page_index = ri.get("page_index", 1)
        image_id = T.make_image_id(ctx.document_id, page_index, img_idx)
        page_id = T.make_page_id(ctx.document_id, page_index)
        images.append(Image(
            image_id=image_id,
            page_id=page_id,
            image_class=ri.get("image_class", "unknown"),
            raw_ocr_text=ri.get("raw_ocr_text", ""),
            clean_ocr_text=ri.get("clean_ocr_text", ""),
            ocr_confidence=ri.get("ocr_confidence"),
            evidence_status=ri.get("evidence_status", "visual_only"),
            confidence=Confidence(overall=ri.get("ocr_confidence", 0.5)),
            traceability=traceability_base,
        ))

    # Detect structural hierarchy from blocks
    sections, articles, clauses = _detect_hierarchy(blocks, ctx.document_id)

    # Build canonical reference IDs for all structural units
    _attach_canonical_ids(sections, articles, clauses, ctx.document_id)

    # Build page refs
    page_ids_seen = sorted({b.page_id for b in blocks})
    pages = [{"page_id": pid, "page_index": i + 1} for i, pid in enumerate(page_ids_seen)]

    # Compute doc-level confidence
    if blocks:
        avg_conf = sum(b.confidence.overall for b in blocks) / len(blocks)
    else:
        avg_conf = 0.0

    document = CanonicalDocument(
        document_id=ctx.document_id,
        schema_version=cfg.schema_version,
        source_filename=ctx.source_path.name,
        mime_type=ctx.get("mime_type", ""),
        file_type=file_type,
        metadata=DocumentMetadata(
            languages=profile_langs,
            document_family=ctx.config.document_type_override or "",
            document_type=ctx.config.document_type_override or "",
        ),
        profile=profile,
        pages=pages,
        blocks=blocks,
        tables=tables,
        images=images,
        sections=sections,
        articles=articles,
        clauses=clauses,
        rendered_outputs=RenderedOutputs(
            markdown=_render_markdown(blocks, tables),
            generation_status="partial" if blocks else "unavailable",
        ),
        confidence=Confidence(overall=avg_conf, extraction=avg_conf),
        traceability=traceability_base,
    )

    ctx.put("document", document)

    return StageOutput(
        stage_name="canonical_structuring",
        status=STATUS_PASS,
        summary=(
            f"Structured {len(blocks)} blocks, {len(tables)} tables, "
            f"{len(sections)} sections, {len(articles)} articles, {len(clauses)} clauses"
        ),
        output_summary={
            "block_count": len(blocks),
            "table_count": len(tables),
            "image_count": len(images),
            "section_count": len(sections),
            "article_count": len(articles),
            "clause_count": len(clauses),
            "structure_detected": bool(sections or articles or clauses),
            "avg_confidence": round(avg_conf, 3),
        },
    )


def _attach_canonical_ids(
    sections: List[Section],
    articles: List[Article],
    clauses: List[Clause],
    document_id: str,
) -> None:
    """
    Attach canonical_id attributes to structure objects.

    Canonical IDs are stored in object.traceability or directly on the object
    via a 'canonical_id' attribute injected into the traceability.repair_methods list.

    Since Pydantic objects may be strict about extra attributes, we store the
    canonical_id on the traceability.repair_methods list as
    "canonical_id:<value>" so it is preserved in the schema and accessible later.

    The CanonicalRefBuilder is run once per document.
    """
    builder = CanonicalRefBuilder(document_id=document_id)

    # Map section_id → canonical_id for parent lookup
    section_canonical: Dict[str, str] = {}
    for section in sections:
        cid = builder.section_ref(section.label or "")
        section_canonical[section.section_id] = cid
        # Store on traceability.repair_methods as "canonical_id:<value>"
        if section.traceability is None:
            section.traceability = Traceability(
                trace_id="", source_hash="", extractor="local_pipeline_v1",
                extractor_version="0.1.0",
            )
        _set_canonical_id(section.traceability, cid)

    # Map article_id → canonical_id for clause parent lookup
    article_canonical: Dict[str, str] = {}
    for article in articles:
        cid = builder.article_ref(article.label or "")
        article_canonical[article.article_id] = cid
        if article.traceability is None:
            article.traceability = Traceability(
                trace_id="", source_hash="", extractor="local_pipeline_v1",
                extractor_version="0.1.0",
            )
        _set_canonical_id(article.traceability, cid)

    for clause in clauses:
        parent_art_canonical = (
            article_canonical.get(clause.parent_article_id or "")
            if clause.parent_article_id else None
        )
        cid = builder.clause_ref(
            clause.label or "",
            parent_canonical=parent_art_canonical,
        )
        if clause.traceability is None:
            clause.traceability = Traceability(
                trace_id="", source_hash="", extractor="local_pipeline_v1",
                extractor_version="0.1.0",
            )
        _set_canonical_id(clause.traceability, cid)


def _set_canonical_id(traceability: Traceability, canonical_id: str) -> None:
    """Store canonical_id in traceability.repair_methods as 'canonical_id:<value>'."""
    # Remove any previous canonical_id entries
    traceability.repair_methods = [
        m for m in traceability.repair_methods
        if not m.startswith("canonical_id:")
    ]
    traceability.repair_methods.append(f"canonical_id:{canonical_id}")


def _get_canonical_id(traceability: Optional[Traceability]) -> Optional[str]:
    """Retrieve canonical_id from traceability.repair_methods."""
    if traceability is None:
        return None
    for m in traceability.repair_methods:
        if m.startswith("canonical_id:"):
            return m[len("canonical_id:"):]
    return None


def _detect_block_type(rb: Dict) -> str:
    """
    Heuristic block type detection.
    Recognises Vietnamese legal structure (Chương, Mục, Điều, Khoản, Điểm)
    and English legal structure (Chapter, Section, Article, Clause, Point).
    """
    style = rb.get("style", "").lower()
    tag = rb.get("tag", "").lower()
    text = (rb.get("raw_text", "") or "").strip()

    # Style / HTML tag signals
    if any(h in style for h in ("heading", "title", "h1", "h2", "h3", "heading 1", "heading 2", "heading 3")):
        return "heading"
    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        return "heading"
    if tag == "li":
        return "list_item"

    # --- Vietnamese hierarchy headings ---
    # Chương I, Chương 1, CHƯƠNG I
    if re.match(r"^\s*chương\s+[\w]+", text, re.IGNORECASE | re.UNICODE):
        return "heading"
    # Mục 1, MỤC 1
    if re.match(r"^\s*mục\s+[\w]+", text, re.IGNORECASE | re.UNICODE):
        return "heading"
    # Điều 5, Điều 5., ĐIỀU 5
    if re.match(r"^\s*điều\s+\d+", text, re.IGNORECASE | re.UNICODE):
        return "heading"
    # Phần I, PHẦN I (Part)
    if re.match(r"^\s*phần\s+[\w]+", text, re.IGNORECASE | re.UNICODE):
        return "heading"

    # --- Vietnamese clause / point markers ---
    # Khoản 1., Khoản 1:
    if re.match(r"^\s*khoản\s+\d+", text, re.IGNORECASE | re.UNICODE):
        return "list_item"
    # Điểm a), Điểm b.
    if re.match(r"^\s*điểm\s+[a-z][\.\):\s]", text, re.IGNORECASE | re.UNICODE):
        return "list_item"

    # --- English hierarchy headings ---
    if re.match(r"^\s*(article|section|chapter|part|schedule|annex|exhibit)\s+[\w]+", text, re.IGNORECASE):
        return "heading"
    if re.match(r"^\s*§\s*\d+", text):
        return "heading"

    # --- Numbered list items (both languages) ---
    # 1. text or 1) text
    if re.match(r"^\s*\d+[\.\)]\s+\S", text):
        return "list_item"
    # (a) text or a) text
    if re.match(r"^\s*[\(\[]?[a-záàảãạăắằẳẵặâấầẩẫậ][\)\]\.]\s+\S", text, re.IGNORECASE | re.UNICODE):
        return "list_item"

    # Short lines ending with colon are often subheadings
    if len(text) < 150 and text.endswith(":") and "\n" not in text:
        return "heading"

    # Citation patterns
    if re.match(r"^\s*\(?\d+\)\s+[A-ZĐÁÀẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬ]", text, re.UNICODE):
        return "list_item"

    return "paragraph"


def _clean_text(text: str) -> str:
    """Mechanical text normalization. Never changes legal wording."""
    text = re.sub(r"\s+", " ", text)
    text = text.strip()
    return text


def _detect_hierarchy(blocks: List[Block], document_id: str):
    """
    Bilingual hierarchy detection from heading blocks.
    Supports Vietnamese legal structure and English legal structure.

    Vietnamese hierarchy:
        Phần / Chương  →  Section (section_kind = "phần" / "chương")
        Mục            →  Section (section_kind = "mục")
        Điều           →  Article
        Khoản / Điểm   →  Clause  (via list_item blocks under a Điều)

    English hierarchy:
        Chapter / Part / Schedule / Annex  →  Section
        Article / Section (numbered)        →  Article
        Clause / Paragraph / Point          →  Clause

    Docs rule: do NOT invent articles or clauses from ambiguous content.
    """
    sections: List[Section] = []
    articles: List[Article] = []
    clauses: List[Clause] = []

    sec_idx = 0
    art_idx = 0
    cls_idx = 0

    current_section_id: Optional[str] = None
    current_article_id: Optional[str] = None

    # ----------------------------------------------------------------
    # Compiled patterns — Vietnamese
    # ----------------------------------------------------------------
    # Phần I, Phần 1, PHẦN I — top-level part
    re_vi_part = re.compile(
        r"^\s*(phần)\s+([\wIVXivx]+)([\.\:\s]+(.*))?$",
        re.IGNORECASE | re.UNICODE,
    )
    # Chương I, Chương 1, CHƯƠNG II
    re_vi_chapter = re.compile(
        r"^\s*(chương)\s+([\wIVXivx]+)([\.\:\s]+(.*))?$",
        re.IGNORECASE | re.UNICODE,
    )
    # Mục 1, Mục 1A
    re_vi_muc = re.compile(
        r"^\s*(mục)\s+(\w+)([\.\:\s]+(.*))?$",
        re.IGNORECASE | re.UNICODE,
    )
    # Điều 5, Điều 5., Điều 5 Tên điều
    re_vi_dieu = re.compile(
        r"^\s*(điều)\s+(\d+[a-zđ]?)([\.\:\s]+(.*))?$",
        re.IGNORECASE | re.UNICODE,
    )
    # Khoản 1, khoản 2
    re_vi_khoan = re.compile(
        r"^\s*(khoản)\s+(\d+)([\.\:\s]+(.*))?$",
        re.IGNORECASE | re.UNICODE,
    )
    # Điểm a, điểm b)
    re_vi_diem = re.compile(
        r"^\s*(điểm)\s+([a-záàảãạăắằẳẵặâấầẩẫậ])([\.\)\:\s]+(.*))?$",
        re.IGNORECASE | re.UNICODE,
    )

    # ----------------------------------------------------------------
    # Compiled patterns — English
    # ----------------------------------------------------------------
    re_en_section = re.compile(
        r"^\s*(chapter|part|title|schedule|annex|exhibit)\s+([\wIVXivx]+)([\.\:\s]+(.*))?$",
        re.IGNORECASE,
    )
    re_en_article = re.compile(
        r"^\s*(article|section|§)\s*(\d+[a-z]?)([\.\:\s]+(.*))?$",
        re.IGNORECASE,
    )
    re_en_clause = re.compile(
        r"^\s*(clause|paragraph|sub-clause|point)\s+(\d+[a-z]?|\([a-z]\))([\.\:\s]+(.*))?$",
        re.IGNORECASE,
    )

    for block in blocks:
        text = block.raw_text.strip()
        if not text:
            continue

        # ---- Vietnamese: Phần / Chương / Mục → Section ----
        for pattern, kind in [
            (re_vi_part, "phần"),
            (re_vi_chapter, "chương"),
            (re_vi_muc, "mục"),
        ]:
            m = pattern.match(text)
            if m:
                section_id = T.make_section_id(document_id, sec_idx)
                sec_idx += 1
                heading_title = (m.group(4) or "").strip() if m.lastindex and m.lastindex >= 4 else ""
                sections.append(Section(
                    section_id=section_id,
                    section_kind=kind,
                    label=text[:120],
                    number=m.group(2),
                    title=heading_title or None,
                    parent_section_id=None,
                    block_ids=[block.block_id],
                    page_refs=[block.page_id],
                    confidence=Confidence(overall=0.90, structure=0.90),
                ))
                current_section_id = section_id
                current_article_id = None
                block.parent_structure_id = section_id
                block.block_type = "heading"
                break
        else:
            # ---- Vietnamese: Điều → Article ----
            m = re_vi_dieu.match(text)
            if m:
                article_id = T.make_article_id(document_id, art_idx)
                art_idx += 1
                heading_title = (m.group(4) or "").strip() if m.lastindex and m.lastindex >= 4 else ""
                articles.append(Article(
                    article_id=article_id,
                    label=text[:120],
                    number=m.group(2),
                    title=heading_title or None,
                    parent_section_id=current_section_id,
                    block_ids=[block.block_id],
                    page_refs=[block.page_id],
                    confidence=Confidence(overall=0.90, structure=0.90),
                ))
                current_article_id = article_id
                block.parent_structure_id = article_id
                block.block_type = "heading"
                continue

            # ---- Vietnamese: Khoản / Điểm → Clause ----
            m_khoan = re_vi_khoan.match(text)
            m_diem = re_vi_diem.match(text)
            if m_khoan or m_diem:
                m = m_khoan or m_diem
                clause_id = T.make_clause_id(document_id, cls_idx)
                cls_idx += 1
                kind = "paragraph" if m_khoan else "point"
                number = m.group(2)
                clauses.append(Clause(
                    clause_id=clause_id,
                    label=text[:120],
                    number=number,
                    clause_kind=kind,
                    parent_article_id=current_article_id,
                    block_ids=[block.block_id],
                    page_refs=[block.page_id],
                    confidence=Confidence(overall=0.85, structure=0.85),
                ))
                block.parent_structure_id = clause_id
                continue

            # ---- English: Chapter / Part / Schedule → Section ----
            m = re_en_section.match(text)
            if m:
                section_id = T.make_section_id(document_id, sec_idx)
                sec_idx += 1
                kind = m.group(1).lower()
                heading_title = (m.group(4) or "").strip() if m.lastindex and m.lastindex >= 4 else ""
                sections.append(Section(
                    section_id=section_id,
                    section_kind=kind,
                    label=text[:120],
                    number=m.group(2),
                    title=heading_title or None,
                    parent_section_id=None,
                    block_ids=[block.block_id],
                    page_refs=[block.page_id],
                    confidence=Confidence(overall=0.90, structure=0.90),
                ))
                current_section_id = section_id
                current_article_id = None
                block.parent_structure_id = section_id
                block.block_type = "heading"
                continue

            # ---- English: Article / Section (numbered) → Article ----
            m = re_en_article.match(text)
            if m:
                article_id = T.make_article_id(document_id, art_idx)
                art_idx += 1
                heading_title = (m.group(4) or "").strip() if m.lastindex and m.lastindex >= 4 else ""
                articles.append(Article(
                    article_id=article_id,
                    label=text[:120],
                    number=m.group(2),
                    title=heading_title or None,
                    parent_section_id=current_section_id,
                    block_ids=[block.block_id],
                    page_refs=[block.page_id],
                    confidence=Confidence(overall=0.88, structure=0.88),
                ))
                current_article_id = article_id
                block.parent_structure_id = article_id
                block.block_type = "heading"
                continue

            # ---- English: Clause / Paragraph → Clause ----
            m = re_en_clause.match(text)
            if m:
                clause_id = T.make_clause_id(document_id, cls_idx)
                cls_idx += 1
                clauses.append(Clause(
                    clause_id=clause_id,
                    label=text[:120],
                    number=m.group(2),
                    clause_kind=m.group(1).lower(),
                    parent_article_id=current_article_id,
                    block_ids=[block.block_id],
                    page_refs=[block.page_id],
                    confidence=Confidence(overall=0.82, structure=0.82),
                ))
                block.parent_structure_id = clause_id
                continue

    return sections, articles, clauses


def _render_markdown(blocks: List[Block], tables: List[Table]) -> str:
    """
    Generate a basic Markdown representation from canonical objects.
    Handles Vietnamese and English headings uniformly.
    Vietnamese diacritics are preserved exactly — never simplified or stripped.
    """
    lines: List[str] = []
    for block in blocks:
        text = block.clean_text or block.raw_text
        if not text:
            continue
        if block.block_type == "heading":
            # Determine heading level from content heuristics
            level = _heading_level(text)
            prefix = "#" * level
            lines.append(f"\n{prefix} {text}\n")
        elif block.block_type == "list_item":
            lines.append(f"- {text}")
        elif block.block_type == "footnote":
            lines.append(f"> {text}")
        else:
            lines.append(text)
    for table in tables:
        if table.markdown:
            lines.append(f"\n{table.markdown}\n")
    return "\n".join(lines)


def _heading_level(text: str) -> int:
    """
    Determine Markdown heading level from text content.
    Vietnamese:  Phần/Chương → h1, Mục → h2, Điều → h3, Khoản/Điểm → h4
    English:     Chapter/Part → h1, Article/Section → h2, Clause → h3
    Default:     h2
    """
    t = text.strip().lower()
    # Level 1 (document part / chapter)
    if re.match(r"^(chương|phần|chapter|part)\s+", t, re.UNICODE):
        return 1
    # Level 2 (section / mục / article)
    if re.match(r"^(mục|section|article|§)\s+", t, re.UNICODE):
        return 2
    # Level 3 (article / điều)
    if re.match(r"^(điều|clause)\s+", t, re.UNICODE):
        return 3
    # Level 4 (khoản / điểm / paragraph)
    if re.match(r"^(khoản|điểm|paragraph|sub-clause)\s+", t, re.UNICODE):
        return 4
    return 2


# ---------------------------------------------------------------------------
# Stage 5: Cleaning and validation
# ---------------------------------------------------------------------------


def stage_cleaning_validation(ctx: StageContext) -> StageOutput:
    """
    Remove obvious garbage, detect missing/duplicate content, flag low-confidence regions.
    Rule: emit warnings, never silently fix legally meaningful text.
    """
    document: CanonicalDocument = ctx.get("document")
    if document is None:
        return StageOutput(
            stage_name="cleaning_validation",
            status=STATUS_FAIL,
            summary="No document object in context",
            errors=["canonical_structuring must run before cleaning_validation"],
        )

    cfg = ctx.config
    warnings: List[str] = []
    issues: List[str] = []

    if cfg.parser_noise_cleanup:
        # Detect duplicated blocks
        seen_texts: Dict[str, str] = {}
        for block in document.blocks:
            normalized = re.sub(r"\s+", " ", block.raw_text.strip().lower())
            if normalized in seen_texts and len(normalized) > 40:
                warnings.append(
                    f"Duplicate block detected: block_id={block.block_id} "
                    f"duplicates {seen_texts[normalized]}"
                )
                block.confidence.degraded = True
                block.confidence.reasons.append("duplicate_content")
            else:
                seen_texts[normalized] = block.block_id

    # Flag low-confidence blocks
    low_conf_blocks = document.low_confidence_blocks(cfg.extraction_confidence_threshold)
    if low_conf_blocks:
        warnings.append(
            f"{len(low_conf_blocks)} block(s) below confidence threshold "
            f"({cfg.extraction_confidence_threshold}): "
            f"{[b.block_id for b in low_conf_blocks[:5]]}"
        )

    # Check coverage
    if document.block_count() == 0:
        issues.append("No blocks extracted. Document may be empty, encrypted, or unsupported.")

    # Validate tables are non-empty
    for table in document.tables:
        if table.row_count == 0:
            warnings.append(f"Table {table.table_id} has 0 rows — may be extraction failure.")

    # Update validation summary
    document.validation = ValidationSummary(
        coverage_score=min(1.0, document.block_count() / max(1, document.profile.page_count * 3)),
        unresolved_issues=issues,
        warnings=warnings,
        validated_at=T.now_iso(),
    )

    status = STATUS_FAIL if issues else (STATUS_WARNING if warnings else STATUS_PASS)
    return StageOutput(
        stage_name="cleaning_validation",
        status=status,
        summary=f"{len(warnings)} warnings, {len(issues)} issues. Coverage: {document.validation.coverage_score:.2f}",
        warnings=warnings,
        errors=issues,
        output_summary={
            "coverage_score": round(document.validation.coverage_score, 3),
            "degraded_blocks": document.degraded_block_count(),
            "low_confidence_blocks": len(low_conf_blocks),
        },
    )


# ---------------------------------------------------------------------------
# Stage 6: Chunking
# ---------------------------------------------------------------------------


def stage_chunking(ctx: StageContext) -> StageOutput:
    """
    Apply chunking strategy from docs/chunking/chunking-strategies.md.
    Preserves legal structure: articles and clauses are never split at the wrong boundary.
    """
    document: CanonicalDocument = ctx.get("document")
    if document is None:
        return StageOutput(
            stage_name="chunking",
            status=STATUS_FAIL,
            summary="No document object in context",
            errors=["canonical_structuring must run before chunking"],
        )

    cfg = ctx.config
    profile = document.profile
    warnings: List[str] = []

    # Choose chunking strategy using the decision tree from the docs
    strategy = _choose_chunk_strategy(profile, document, cfg)
    decision = ChunkingDecision(
        document_id=ctx.document_id,
        strategy=strategy,
        secondary_rules=_secondary_rules(profile, document),
        fallback_strategy=_fallback_strategy(strategy),
        routing_signals={
            "page_count": profile.page_count,
            "is_long_document": profile.is_long_document,
            "has_tables": profile.has_tables,
            "table_density": profile.table_density,
            "image_density": profile.image_density,
        },
        decision_confidence=0.85,
        reason_codes=[f"strategy={strategy}"],
        created_at=T.now_iso(),
    )

    ctx.logger.decision(
        stage="chunking",
        decision_type="chunking_strategy",
        value=strategy,
        reason=f"pages={profile.page_count}, long={profile.is_long_document}, tables={profile.has_tables}",
    )

    # Build chunks
    chunks: List[Chunk] = []
    chunk_idx = 0

    # --- Text chunks: group by structural unit (article → group of blocks) ---
    if document.articles:
        for article in document.articles:
            # Collect body blocks: those explicitly listed + those parented to this article
            body_blocks = [
                b for b in document.blocks
                if b.parent_structure_id == article.article_id
                or b.block_id in article.block_ids
            ]
            body_text = "\n\n".join(
                (b.clean_text or b.raw_text)
                for b in body_blocks
                if (b.clean_text or b.raw_text)
            )
            structure_path = _build_path(article, document)

            # Chunk content: include structure path as context header so keyword
            # searches on hierarchy terms ("điều", "chương", "article", etc.) always hit.
            path_header = " › ".join(structure_path)
            content = f"{path_header}\n\n## {article.label}\n\n{body_text}" if body_text \
                else f"{path_header}\n\n## {article.label}"

            # Multilingual: canonical refs and language metadata
            chunk_canonical_refs = canonical_refs_for_chunk(structure_path, content)
            chunk_lang = _get_chunk_language(document)
            hierarchy_path_str = " › ".join(structure_path)

            chunk = Chunk(
                chunk_id=T.make_chunk_id(ctx.document_id, chunk_idx),
                document_id=ctx.document_id,
                chunk_type="text",
                content_format="markdown",
                content=content,
                structure_path=structure_path,
                page_refs=article.page_refs,
                block_refs=[b.block_id for b in body_blocks] or article.block_ids,
                citations=article.citations,
                token_estimate=len(content) // 4,
                confidence=article.confidence.overall,
                degraded=article.confidence.degraded,
                jurisdiction=document.metadata.jurisdiction,
                version_label=document.metadata.version_label or "",
                document_type=document.metadata.document_type,
                language=chunk_lang,
                canonical_refs=chunk_canonical_refs,
                hierarchy_path=hierarchy_path_str,
            )
            chunks.append(chunk)
            chunk_idx += 1
    else:
        # No detected hierarchy: chunk by paragraph groups (fallback)
        CHUNK_BLOCK_SIZE = 10 if not profile.is_long_document else 20
        for i in range(0, max(1, len(document.blocks)), CHUNK_BLOCK_SIZE):
            group = document.blocks[i: i + CHUNK_BLOCK_SIZE]
            if not group:
                continue
            content = "\n\n".join(
                (b.clean_text or b.raw_text)
                for b in group
                if (b.clean_text or b.raw_text)
            )
            avg_conf = sum(b.confidence.overall for b in group) / len(group)
            chunk_canonical_refs = canonical_refs_for_chunk([], content)
            chunk_lang = _get_chunk_language(document)

            chunk = Chunk(
                chunk_id=T.make_chunk_id(ctx.document_id, chunk_idx),
                document_id=ctx.document_id,
                chunk_type="text",
                content_format="markdown",
                content=content,
                structure_path=[],
                page_refs=list({b.page_id for b in group}),
                block_refs=[b.block_id for b in group],
                token_estimate=len(content) // 4,
                confidence=avg_conf,
                degraded=avg_conf < cfg.chunk_authority_threshold,
                jurisdiction=document.metadata.jurisdiction,
                version_label=document.metadata.version_label or "",
                document_type=document.metadata.document_type,
                language=chunk_lang,
                canonical_refs=chunk_canonical_refs,
                hierarchy_path="",
            )
            chunks.append(chunk)
            chunk_idx += 1

    # --- Table chunks: one chunk per table ---
    for table in document.tables:
        content = table.projection_text or table.markdown or "[table]"
        chunk_canonical_refs = canonical_refs_for_chunk([], content)
        chunk_lang = _get_chunk_language(document)

        chunk = Chunk(
            chunk_id=T.make_chunk_id(ctx.document_id, chunk_idx),
            document_id=ctx.document_id,
            chunk_type="table",
            content_format="markdown",
            content=content,
            structure_path=[],
            page_refs=table.page_refs,
            table_refs=[table.table_id],
            token_estimate=len(content) // 4,
            confidence=table.confidence.overall,
            degraded=table.confidence.degraded,
            jurisdiction=document.metadata.jurisdiction,
            version_label=document.metadata.version_label or "",
            document_type=document.metadata.document_type,
            language=chunk_lang,
            canonical_refs=chunk_canonical_refs,
            hierarchy_path="",
        )
        chunks.append(chunk)
        chunk_idx += 1

    degraded_count = sum(1 for c in chunks if c.degraded)
    chunk_set = ChunkSet(
        document_id=ctx.document_id,
        chunks=chunks,
        chunking_decision=decision,
        total_chunks=len(chunks),
        degraded_chunks=degraded_count,
        strategy_used=strategy,
        created_at=T.now_iso(),
    )
    ctx.put("chunk_set", chunk_set)

    return StageOutput(
        stage_name="chunking",
        status=STATUS_WARNING if warnings else STATUS_PASS,
        summary=f"{len(chunks)} chunks ({degraded_count} degraded) using strategy={strategy}",
        warnings=warnings,
        output_summary={
            "total_chunks": len(chunks),
            "text_chunks": len(chunk_set.text_chunks()),
            "table_chunks": len(chunk_set.table_chunks()),
            "degraded_chunks": degraded_count,
            "strategy": strategy,
        },
    )


def _get_chunk_language(document: CanonicalDocument) -> str:
    """
    Determine the primary language for chunks in a document.

    Uses the formal language_detection result if available (confidence >= 0.5),
    otherwise falls back to the first entry in profile.languages.
    """
    if document.profile.language_detection:
        ld = document.profile.language_detection
        if ld.confidence >= 0.5 and ld.language not in ("unknown",):
            return ld.language
    langs = document.profile.languages or document.metadata.languages
    return langs[0] if langs else "en"


def _choose_chunk_strategy(
    profile: DocumentProfile, document: CanonicalDocument, cfg: PipelineConfig
) -> str:
    """Implement the chunking decision tree from docs."""
    structure_score = 0.8 if document.has_structure() else 0.3
    if structure_score >= 0.7 and profile.table_density <= 0.2 and profile.image_density <= 0.2:
        return "long_local_structural" if profile.is_long_document else "structural"
    if structure_score >= 0.5 and profile.is_long_document:
        return "legal_aware"
    if profile.table_density >= 0.4 and profile.layout_complexity_score <= 0.5:
        return "table_aware"
    if profile.image_density >= 0.3 or (profile.scan_quality_score < 0.5):
        return "conservative_fallback"
    if profile.layout_complexity_score >= 0.7:
        return "mixed_group"
    return "semantic"


def _secondary_rules(profile: DocumentProfile, document: CanonicalDocument) -> List[str]:
    rules = []
    if document.tables:
        rules.append("enforce_table_header_preservation")
    if document.images:
        rules.append("create_evidence_sibling_chunks")
    return rules


def _fallback_strategy(strategy: str) -> str:
    fallback_map = {
        "structural": "legal_aware",
        "long_local_structural": "legal_aware",
        "legal_aware": "semantic",
        "table_aware": "mixed_group",
        "mixed_group": "conservative_fallback",
        "semantic": "conservative_fallback",
        "conservative_fallback": "conservative_fallback",
    }
    return fallback_map.get(strategy, "conservative_fallback")


def _build_path(article: Article, document: CanonicalDocument) -> List[str]:
    path = []
    if article.parent_section_id:
        for sec in document.sections:
            if sec.section_id == article.parent_section_id:
                path.append(sec.label or sec.title or sec.section_id)
                break
    path.append(article.label or article.article_id)
    return path


# ---------------------------------------------------------------------------
# Stage 7: Graph building
# ---------------------------------------------------------------------------


def stage_graph_building(ctx: StageContext) -> StageOutput:
    """
    Build graph nodes and edges from the canonical document.
    Conservative: structural edges first, semantic edges only with evidence.
    """
    if not ctx.config.enable_graph_building:
        return StageOutput(
            stage_name="graph_building",
            status="skipped",
            summary="Graph building disabled by config.",
        )

    document: CanonicalDocument = ctx.get("document")
    chunk_set: ChunkSet = ctx.get("chunk_set")
    if document is None:
        return StageOutput(
            stage_name="graph_building",
            status=STATUS_FAIL,
            summary="No document in context",
            errors=["canonical_structuring must run before graph_building"],
        )

    nodes: List[GraphNode] = []
    edges: List[GraphEdge] = []
    job_id = str(uuid.uuid4())

    from src.schemas.graph import GraphNodeProvenance

    def _make_node(node_type: str, label: str, source_ref: str, attrs: Dict = {}) -> GraphNode:
        return GraphNode(
            node_id=T.make_node_id(ctx.document_id, node_type, source_ref),
            node_type=node_type,
            label=label[:120],
            document_scope=ctx.document_id,
            source_trace=GraphNodeProvenance(
                document_id=ctx.document_id,
                source_anchor=source_ref,
                extraction_method="local",
                processing_version=ctx.config.processing_version,
                created_at=T.now_iso(),
            ),
            confidence=0.90,
            attributes=attrs,
        )

    def _make_edge(from_id: str, edge_type: str, to_id: str, confidence: float = 0.9) -> GraphEdge:
        return GraphEdge(
            edge_id=T.make_edge_id(from_id, edge_type, to_id),
            edge_type=edge_type,
            from_node_id=from_id,
            to_node_id=to_id,
            confidence=confidence,
            provenance=ctx.document_id,
            method="structural",
        )

    # Document root node
    doc_node = _make_node("Document", document.source_filename, document.document_id)
    nodes.append(doc_node)

    # Section nodes
    section_node_ids: Dict[str, str] = {}
    for section in document.sections:
        cid = _get_canonical_id(section.traceability)
        attrs: Dict = {}
        if cid:
            attrs["canonical_refs"] = [cid]
        n = _make_node("Section", section.label or section.section_id, section.section_id, attrs=attrs)
        nodes.append(n)
        section_node_ids[section.section_id] = n.node_id
        edges.append(_make_edge(doc_node.node_id, "CONTAINS", n.node_id))

    # Article nodes
    article_node_ids: Dict[str, str] = {}
    for article in document.articles:
        cid = _get_canonical_id(article.traceability)
        attrs = {}
        if cid:
            attrs["canonical_refs"] = [cid]
        n = _make_node("Article", article.label or article.article_id, article.article_id, attrs=attrs)
        nodes.append(n)
        article_node_ids[article.article_id] = n.node_id
        if article.parent_section_id and article.parent_section_id in section_node_ids:
            edges.append(_make_edge(section_node_ids[article.parent_section_id], "CONTAINS", n.node_id))
        else:
            edges.append(_make_edge(doc_node.node_id, "CONTAINS", n.node_id))

    # Table nodes
    for table in document.tables:
        n = _make_node("Table", table.title or table.table_id, table.table_id)
        nodes.append(n)
        edges.append(_make_edge(doc_node.node_id, "HAS_TABLE", n.node_id))

    # Image evidence nodes
    for image in document.images:
        n = _make_node("ImageEvidence", image.image_class or image.image_id, image.image_id)
        nodes.append(n)
        edges.append(_make_edge(doc_node.node_id, "HAS_IMAGE", n.node_id))

    # Chunk nodes — include canonical_refs in attributes so alias edges can be built
    if chunk_set:
        for chunk in chunk_set.chunks:
            chunk_attrs: Dict = {"chunk_type": chunk.chunk_type}
            chunk_canonical = getattr(chunk, "canonical_refs", None) or []
            if chunk_canonical:
                chunk_attrs["canonical_refs"] = chunk_canonical
            n = _make_node("Chunk", chunk.chunk_id, chunk.chunk_id, attrs=chunk_attrs)
            n.confidence = chunk.confidence
            nodes.append(n)
            edges.append(_make_edge(doc_node.node_id, "DERIVED_TO_CHUNK", n.node_id, confidence=0.95))

    graph = GraphSubgraph(
        document_id=ctx.document_id,
        nodes=nodes,
        edges=edges,
    )

    # Enrich graph with multilingual ALIAS_OF edges
    graph = add_alias_edges(graph)

    graph.build_job = GraphBuildJob(
        job_id=job_id,
        document_id=ctx.document_id,
        processing_version=ctx.config.processing_version,
        created_nodes=[n.node_id for n in graph.nodes],
        created_edges=[e.edge_id for e in graph.edges],
        validation_status="passed",
        nodes_by_type=graph_nodes_by_type(graph.nodes),
        edges_by_type=graph_edges_by_type(graph.edges),
    )
    ctx.put("graph", graph)

    alias_edge_count = sum(1 for e in graph.edges if e.edge_type == "ALIAS_OF")
    return StageOutput(
        stage_name="graph_building",
        status=STATUS_PASS,
        summary=(
            f"{graph.node_count()} nodes, {graph.edge_count()} edges "
            f"({alias_edge_count} ALIAS_OF)"
        ),
        output_summary={
            "total_nodes": graph.node_count(),
            "total_edges": graph.edge_count(),
            "alias_edges": alias_edge_count,
            "nodes_by_type": graph.nodes_by_type(),
            "edges_by_type": graph.edges_by_type(),
            "structural_edges": graph.structural_edge_count(),
        },
    )


def graph_nodes_by_type(nodes: List[GraphNode]) -> Dict[str, int]:
    counts: Dict[str, int] = {}
    for n in nodes:
        counts[n.node_type] = counts.get(n.node_type, 0) + 1
    return counts


def graph_edges_by_type(edges: List[GraphEdge]) -> Dict[str, int]:
    counts: Dict[str, int] = {}
    for e in edges:
        counts[e.edge_type] = counts.get(e.edge_type, 0) + 1
    return counts


# ---------------------------------------------------------------------------
# Stage 8: Retrieval smoke test
# ---------------------------------------------------------------------------


def stage_retrieval_smoke_test(ctx: StageContext) -> StageOutput:
    """
    Multilingual retrieval smoke test over chunks using the RetrievalEngine.

    Upgrades over the previous keyword-only test:
    - Uses RetrievalEngine with canonical ref + alias expansion
    - Generates language-appropriate primary queries
    - Also tests cross-language queries (English queries against a VI doc, and vice versa)
    - Logs retrieval diagnostics via ctx.logger

    Design rules:
    - queries is always a LOCAL copy; config.smoke_test_queries is never mutated.
    - All-queries-fail → STATUS_WARNING (never STATUS_FAIL from smoke test alone).
    - Cross-language hit rate is computed separately and logged but does not gate status.
    """
    if not ctx.config.enable_retrieval_smoke_test:
        return StageOutput(
            stage_name="retrieval_smoke_test",
            status="skipped",
            summary="Retrieval smoke test disabled by config.",
        )

    chunk_set: ChunkSet = ctx.get("chunk_set")
    if chunk_set is None or not chunk_set.chunks:
        return StageOutput(
            stage_name="retrieval_smoke_test",
            status=STATUS_WARNING,
            summary="No chunks available for smoke test.",
            warnings=["chunk_set is empty; no retrieval test possible"],
        )

    document: CanonicalDocument = ctx.get("document")

    # Build the multilingual retrieval engine (debug=True for logging)
    engine = RetrievalEngine(chunk_set, debug=True)

    # ----------------------------------------------------------------
    # Step 1: Determine document language
    # ----------------------------------------------------------------
    doc_lang = "en"
    if document:
        if document.profile.language_detection:
            ld = document.profile.language_detection
            if ld.confidence >= 0.5 and ld.language not in ("unknown",):
                doc_lang = ld.language
        elif document.profile.languages:
            doc_lang = document.profile.languages[0]

    # ----------------------------------------------------------------
    # Step 2: Build primary queries (in document language)
    # Always work with a LOCAL copy — never mutate config
    # ----------------------------------------------------------------
    primary_queries: List[str] = list(ctx.config.smoke_test_queries)  # copy, not reference

    if not primary_queries:
        if doc_lang == "vi":
            primary_queries = [
                "điều", "khoản", "hợp đồng", "bên", "nghĩa vụ",
                "thanh toán", "chương", "mục",
            ]
        elif doc_lang == "mixed":
            primary_queries = [
                "điều", "article", "khoản", "clause", "hợp đồng", "agreement",
            ]
        else:  # English default
            primary_queries = [
                "article", "clause", "agreement", "obligation", "payment",
                "chapter", "section",
            ]

    # ----------------------------------------------------------------
    # Step 3: Build cross-language queries (opposite language)
    # These test canonical ref matching across language boundaries
    # ----------------------------------------------------------------
    cross_lang_queries: List[str] = []
    if document and document.articles:
        if doc_lang == "vi":
            # Test English canonical form of the first few article labels
            for art in document.articles[:3]:
                if art.number:
                    cross_lang_queries.append(f"Article {art.number}")
        elif doc_lang == "en":
            # Test Vietnamese canonical form of the first few article labels
            for art in document.articles[:3]:
                if art.number:
                    cross_lang_queries.append(f"Điều {art.number}")

    # ----------------------------------------------------------------
    # Step 4: Run primary retrieval
    # ----------------------------------------------------------------
    primary_results: List[int] = []
    failed_primary: List[str] = []

    for query in primary_queries[:10]:
        hits = engine.search(query)
        primary_results.append(len(hits))
        if not hits:
            failed_primary.append(query)

        # Log retrieval diagnostics
        if engine.last_debug_log:
            log = engine.last_debug_log
            ctx.logger.info(
                f"Smoke query: {query!r} → {len(hits)} hits "
                f"(canonical={log.canonical_hits}, alias={log.alias_hits}, "
                f"keyword={log.keyword_hits})",
                stage="retrieval_smoke_test",
            )

    avg_hits = sum(primary_results) / max(1, len(primary_results))
    hit_rate = (len(primary_queries) - len(failed_primary)) / max(1, len(primary_queries))

    # ----------------------------------------------------------------
    # Step 5: Run cross-language retrieval (advisory, not gating)
    # ----------------------------------------------------------------
    cross_lang_results: Dict[str, List] = {}
    for query in cross_lang_queries[:5]:
        hits = engine.search(query)
        cross_lang_results[query] = hits
        if engine.last_debug_log:
            log = engine.last_debug_log
            ctx.logger.info(
                f"Cross-lang query: {query!r} → {len(hits)} hits "
                f"(canonical={log.canonical_hits}, alias={log.alias_hits})",
                stage="retrieval_smoke_test",
            )

    cross_lang_hit_rate = 0.0
    if cross_lang_queries:
        cross_hits = sum(1 for r in cross_lang_results.values() if r)
        cross_lang_hit_rate = cross_hits / len(cross_lang_queries)

    # ----------------------------------------------------------------
    # Step 6: Determine status (primary queries only gate status)
    # ----------------------------------------------------------------
    if hit_rate >= 0.6:
        status = STATUS_PASS
    else:
        status = STATUS_WARNING

    # Build warnings
    warnings: List[str] = [f"No results for query: '{q}'" for q in failed_primary]
    if cross_lang_queries and cross_lang_hit_rate < 0.5:
        failed_cross = [q for q, r in cross_lang_results.items() if not r]
        for q in failed_cross:
            warnings.append(f"Cross-language query missed: '{q}'")

    summary = (
        f"Tested {len(primary_queries)} queries ({doc_lang}) — "
        f"{len(primary_queries) - len(failed_primary)} hit, {len(failed_primary)} missed "
        f"(hit rate {hit_rate:.0%}, avg {avg_hits:.1f} chunks/query)"
    )
    if cross_lang_queries:
        summary += (
            f" | Cross-lang: {int(cross_lang_hit_rate * 100)}% "
            f"({len(cross_lang_queries)} queries)"
        )

    ctx.put("cross_lang_retrieval_results", cross_lang_results)

    return StageOutput(
        stage_name="retrieval_smoke_test",
        status=status,
        summary=summary,
        warnings=warnings,
        output_summary={
            "doc_language": doc_lang,
            "queries_tested": len(primary_queries),
            "queries_with_results": len(primary_queries) - len(failed_primary),
            "queries_failed": failed_primary,
            "avg_result_count": round(avg_hits, 1),
            "hit_rate": round(hit_rate, 3),
            "cross_lang_queries_tested": len(cross_lang_queries),
            "cross_lang_hit_rate": round(cross_lang_hit_rate, 3),
        },
    )


def _keyword_search(query: str, chunk_set: ChunkSet) -> List[Chunk]:
    """Simple case-insensitive keyword search over chunk content. Used as fallback."""
    query_lower = query.lower()
    return [c for c in chunk_set.chunks if query_lower in c.content.lower()]
