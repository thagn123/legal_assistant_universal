"""
Stage 1: Input loading
Stage 2: Document profiling

Heuristic document profiling — local-first, no AI, deterministic.
Language support: Vietnamese (vi) and English (en) legal hierarchies.
"""

from __future__ import annotations

import mimetypes
import re
from pathlib import Path
from typing import List, Optional

from src.config import PipelineConfig, ExtractionStrategy
from src.pipeline.interfaces import StageContext, StageOutput
from src.schemas.document import DocumentProfile, LanguageDetection
from src.schemas.evaluation import STATUS_FAIL, STATUS_PASS, STATUS_WARNING
from src.utils import trace as T
from src.retrieval.language_detector import detect_language as _formal_detect_language


# ---------------------------------------------------------------------------
# Stage 1: Input loading
# ---------------------------------------------------------------------------


def stage_input_loading(ctx: StageContext) -> StageOutput:
    """
    Validate the source file, detect file type, collect basic metadata.
    Computes source_hash for provenance.
    """
    path = ctx.source_path
    warnings: List[str] = []
    errors: List[str] = []

    if not path.exists():
        return StageOutput(
            stage_name="input_loading",
            status=STATUS_FAIL,
            summary=f"Source file not found: {path}",
            errors=[f"File not found: {path}"],
        )

    # Compute source hash
    source_hash = T.compute_source_hash(path)

    # Detect MIME type and file extension
    mime, _ = mimetypes.guess_type(str(path))
    suffix = path.suffix.lower().lstrip(".")

    file_type_map = {
        "pdf": "pdf",
        "docx": "docx",
        "doc": "doc",       # legacy Word — handled by its own extractor
        "html": "html",
        "htm": "html",
        "png": "image",
        "jpg": "image",
        "jpeg": "image",
        "tiff": "image",
        "tif": "image",
        "bmp": "image",
        "webp": "image",
    }
    file_type = file_type_map.get(suffix, "unknown")

    if file_type == "unknown":
        warnings.append(f"Unknown file type for extension '.{suffix}'; will attempt text extraction.")

    file_size_bytes = path.stat().st_size
    ctx.put("source_hash", source_hash)
    ctx.put("file_type", file_type)
    ctx.put("mime_type", mime or "application/octet-stream")
    ctx.put("file_size_bytes", file_size_bytes)

    ctx.logger.info(
        f"Loaded: {path.name} | type={file_type} | size={file_size_bytes:,}B | hash={source_hash[:12]}",
        stage="input_loading",
    )

    return StageOutput(
        stage_name="input_loading",
        status=STATUS_WARNING if warnings else STATUS_PASS,
        summary=f"Loaded '{path.name}' (type={file_type}, size={file_size_bytes:,} bytes)",
        warnings=warnings,
        errors=errors,
        output_summary={
            "filename": path.name,
            "file_type": file_type,
            "file_size_bytes": file_size_bytes,
            "source_hash": source_hash,
        },
    )


# ---------------------------------------------------------------------------
# Stage 2: Document profiling
# ---------------------------------------------------------------------------


def stage_document_profiling(ctx: StageContext) -> StageOutput:
    """
    Profile the document to determine extraction strategy.
    Local-first: uses file introspection without AI.
    """
    path = ctx.source_path
    file_type = ctx.get("file_type", "unknown")
    cfg = ctx.config
    warnings: List[str] = []

    profile = DocumentProfile(file_type=file_type)

    try:
        if file_type == "pdf":
            profile = _profile_pdf(path, cfg, warnings)
        elif file_type == "docx":
            profile = _profile_docx(path, cfg, warnings)
        elif file_type == "doc":
            profile = _profile_doc(path, cfg, warnings)
        elif file_type == "html":
            profile = _profile_html(path, cfg, warnings)
        elif file_type == "image":
            profile = _profile_image(path, cfg, warnings)
        else:
            warnings.append(f"Cannot profile unknown file type '{file_type}'; using defaults.")
            profile.extraction_strategy = ExtractionStrategy.SIMPLE_LOCAL
    except Exception as exc:
        warnings.append(f"Profiling error: {exc}. Using conservative defaults.")
        profile.extraction_strategy = ExtractionStrategy.SIMPLE_LOCAL

    ctx.put("profile", profile)
    ctx.logger.decision(
        stage="document_profiling",
        decision_type="extraction_strategy",
        value=profile.extraction_strategy,
        reason=(
            f"text_layer={profile.text_layer_coverage:.2f}, "
            f"scan_quality={profile.scan_quality_score:.2f}, "
            f"layout_complexity={profile.layout_complexity_score:.2f}, "
            f"pages={profile.page_count}, "
            f"long_doc={profile.is_long_document}"
        ),
        thresholds={
            "strong_text_layer": cfg.strong_text_layer_threshold,
            "weak_text_layer": cfg.weak_text_layer_threshold,
            "long_document_pages": cfg.long_document_page_threshold,
        },
    )

    return StageOutput(
        stage_name="document_profiling",
        status=STATUS_WARNING if warnings else STATUS_PASS,
        summary=f"Strategy: {profile.extraction_strategy} | pages={profile.page_count} | tables={profile.has_tables}",
        warnings=warnings,
        output_summary={
            "extraction_strategy": profile.extraction_strategy,
            "page_count": profile.page_count,
            "text_layer_coverage": round(profile.text_layer_coverage, 3),
            "is_long_document": profile.is_long_document,
            "has_tables": profile.has_tables,
            "has_images": profile.has_images,
        },
    )


# ---------------------------------------------------------------------------
# Profiling helpers
# ---------------------------------------------------------------------------


def _profile_pdf(path: Path, cfg: PipelineConfig, warnings: List[str]) -> DocumentProfile:
    """Profile a PDF using pdfminer or pypdf if available."""
    profile = DocumentProfile(file_type="pdf")
    try:
        import pdfminer.high_level as pdfminer
        from pdfminer.layout import LAParams

        text = pdfminer.extract_text(str(path), laparams=LAParams())
        page_count = max(1, text.count("\x0c") + 1)
        text_length = len(text.strip())

        profile.page_count = page_count
        profile.text_layer_coverage = min(1.0, text_length / max(1, page_count * 500))
        profile.estimated_token_count = text_length // 4  # rough token estimate

        # Heuristic table detection
        profile.has_tables = bool(re.search(r"\|.*\||\t.*\t", text))
        profile.is_long_document = (
            page_count >= cfg.long_document_page_threshold
            or profile.estimated_token_count >= cfg.long_document_token_threshold
        )
        profile.is_scan_heavy = profile.text_layer_coverage < cfg.weak_text_layer_threshold
        # Language detection (heuristic + formal)
        profile.languages = _detect_languages(text)
        _attach_language_detection(profile, text)

    except ImportError:
        warnings.append("pdfminer not installed; using minimal PDF profiling.")
        profile.page_count = 1
        profile.text_layer_coverage = 0.5
        profile.extraction_strategy = ExtractionStrategy.SIMPLE_LOCAL
        return profile

    # Route extraction strategy
    profile.extraction_strategy = _choose_strategy(profile, cfg)
    return profile


def _profile_docx(path: Path, cfg: PipelineConfig, warnings: List[str]) -> DocumentProfile:
    profile = DocumentProfile(file_type="docx")
    try:
        import docx  # python-docx

        doc = docx.Document(str(path))
        para_text = " ".join(p.text for p in doc.paragraphs)
        profile.page_count = max(1, len(para_text) // 3000)
        profile.text_layer_coverage = 1.0
        profile.has_tables = len(doc.tables) > 0
        profile.estimated_token_count = len(para_text) // 4
        profile.is_long_document = profile.estimated_token_count >= cfg.long_document_token_threshold
        profile.languages = _detect_languages(para_text)
        _attach_language_detection(profile, para_text)
    except ImportError:
        warnings.append("python-docx not installed; using fallback DOCX profiling.")
        profile.text_layer_coverage = 1.0

    profile.extraction_strategy = (
        ExtractionStrategy.SIMPLE_LOCAL if not profile.is_long_document
        else ExtractionStrategy.LONG_LOCAL
    )
    return profile


def _profile_doc(path: Path, cfg: PipelineConfig, warnings: List[str]) -> DocumentProfile:
    """
    Profile a legacy .doc file.
    Uses _read_doc_text which tries win32com → LibreOffice → docx2txt → antiword → python-docx.
    """
    # Import here to avoid circular dependency at module load time
    from src.pipeline.extractor import _read_doc_text

    profile = DocumentProfile(file_type="doc")
    text = _read_doc_text(path, warnings)
    if text:
        profile.text_layer_coverage = 1.0
        profile.estimated_token_count = len(text) // 4
        profile.has_tables = False  # docx2txt loses table topology; mark as missing
        profile.is_long_document = profile.estimated_token_count >= cfg.long_document_token_threshold
        # Detect Vietnamese or English
        profile.languages = _detect_languages(text)
        _attach_language_detection(profile, text)
    else:
        warnings.append(".doc file produced no text. File may be corrupt or binary-only.")
        profile.text_layer_coverage = 0.0
    profile.extraction_strategy = (
        ExtractionStrategy.SIMPLE_LOCAL if not profile.is_long_document
        else ExtractionStrategy.LONG_LOCAL
    )
    return profile


def _profile_html(path: Path, cfg: PipelineConfig, warnings: List[str]) -> DocumentProfile:
    profile = DocumentProfile(file_type="html")
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        profile.text_layer_coverage = 1.0
        # HTML files are single logical "page" — use token count for page estimate
        char_len = len(text)
        profile.estimated_token_count = char_len // 4
        profile.page_count = max(1, profile.estimated_token_count // 2000)  # ~2000 tokens/page
        profile.has_tables = "<table" in text.lower()
        profile.has_images = "<img" in text.lower()
        profile.is_long_document = profile.estimated_token_count >= cfg.long_document_token_threshold
        profile.languages = _detect_languages(text)
        _attach_language_detection(profile, text)
    except Exception as exc:
        warnings.append(f"HTML read error: {exc}")
        profile.page_count = 1
    profile.extraction_strategy = ExtractionStrategy.SIMPLE_LOCAL
    return profile


def _profile_image(path: Path, cfg: PipelineConfig, warnings: List[str]) -> DocumentProfile:
    profile = DocumentProfile(file_type="image")
    profile.text_layer_coverage = 0.0
    profile.is_scan_heavy = True
    profile.image_density = 1.0
    if not cfg.enable_ocr:
        warnings.append("Image file detected but OCR is disabled; text extraction will be empty.")
    profile.extraction_strategy = ExtractionStrategy.SCAN_RECOVERY
    profile.page_count = 1
    return profile


def _attach_language_detection(profile: DocumentProfile, text: str) -> None:
    """
    Run the formal language detector and attach the result to the document profile.
    Uses the first 6000 chars for performance; updates profile.language_detection in-place.
    Also syncs profile.languages from the detection result when confident (>= 0.6).
    """
    result = _formal_detect_language(text[:6000])
    profile.language_detection = LanguageDetection(
        language=result.language,
        confidence=result.confidence,
        jurisdiction=result.jurisdiction,
        script=result.script,
        signals=result.signals[:15],
    )
    # Sync profile.languages if detection is confident
    if result.confidence >= 0.6 and result.language not in ("unknown", "mixed"):
        if result.language not in profile.languages:
            profile.languages = [result.language] + [
                l for l in profile.languages if l != result.language
            ]
    elif result.language == "mixed":
        # Ensure both vi and en are present for mixed documents
        for lang in ("vi", "en"):
            if lang not in profile.languages:
                profile.languages.append(lang)


def _detect_languages(text: str) -> List[str]:
    """
    Heuristic language detection for Vietnamese and English.
    Vietnamese is detected by the presence of diacritical characters unique to it.
    Returns a list of detected language codes (e.g. ["vi"], ["en"], ["vi", "en"]).
    """
    # Vietnamese-specific diacritics (tonal marks and vowel modifiers)
    vi_pattern = re.compile(
        r"[àáâãèéêìíòóôõùúýăđơưạảấầẩẫậắằẳẵặẹẻẽếềểễệỉịọỏốồổỗộớờởỡợụủứừửữựỳỷỹỵ"
        r"ÀÁÂÃÈÉÊÌÍÒÓÔÕÙÚÝĂĐƠƯẠẢẤẦẨẪẬẮẰẲẴẶẸẺẼẾỀỂỄỆỈỊỌỎỐỒỔỖỘỚỜỞỠỢỤỦỨỪỬỮỰỲỶỸỴ]",
        re.UNICODE,
    )
    # Vietnamese legal keywords
    vi_keywords = re.compile(
        r"\b(điều|khoản|điểm|chương|mục|nghị định|thông tư|luật|bộ luật|quyết định|nghị quyết)\b",
        re.IGNORECASE | re.UNICODE,
    )
    # English legal keywords
    en_keywords = re.compile(
        r"\b(article|clause|section|chapter|whereas|herein|agreement|contract|whereas|shall|obligation)\b",
        re.IGNORECASE,
    )

    langs = []
    sample = text[:5000]  # check first 5000 chars for speed

    has_vi_chars = bool(vi_pattern.search(sample))
    has_vi_keywords = bool(vi_keywords.search(sample))
    has_en_keywords = bool(en_keywords.search(sample))

    if has_vi_chars or has_vi_keywords:
        langs.append("vi")
    if has_en_keywords:
        langs.append("en")
    if not langs:
        langs.append("en")  # safe default
    return langs


def _choose_strategy(profile: DocumentProfile, cfg: PipelineConfig) -> str:
    """
    Routing logic from docs/parsing/document-intelligence-pipeline.md decision tree.
    """
    if profile.file_type in ("docx", "html"):
        return (
            ExtractionStrategy.SIMPLE_LOCAL if not profile.is_long_document
            else ExtractionStrategy.LONG_LOCAL
        )
    if profile.text_layer_coverage >= cfg.strong_text_layer_threshold:
        if profile.is_long_document:
            return ExtractionStrategy.LONG_LOCAL
        return ExtractionStrategy.SIMPLE_LOCAL
    if profile.text_layer_coverage < cfg.weak_text_layer_threshold:
        return ExtractionStrategy.SCAN_RECOVERY
    return ExtractionStrategy.HYBRID_REGION_PRECISION
