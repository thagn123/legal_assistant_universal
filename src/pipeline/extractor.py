"""
Stage 3: Extraction

Local-first text, table, and image extraction from supported file formats.
AI repair is gated by config.enable_ai_repair and never runs by default.

File format support:
- .pdf  — pdfminer (local text layer)
- .docx — python-docx
- .doc  — win32com → LibreOffice → docx2txt → antiword → python-docx
- .html — beautifulsoup4
- image — pytesseract OCR
"""

from __future__ import annotations

import re
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from src.config import PipelineConfig
from src.pipeline.interfaces import StageContext, StageOutput
from src.schemas.document import DocumentProfile
from src.schemas.evaluation import STATUS_FAIL, STATUS_PASS, STATUS_WARNING


# ---------------------------------------------------------------------------
# Stage 3: Extraction
# ---------------------------------------------------------------------------


def stage_extraction(ctx: StageContext) -> StageOutput:
    """
    Extract text, tables, and images from the source document.
    Local-first; AI repair only if explicitly enabled.
    """
    path = ctx.source_path
    profile: DocumentProfile = ctx.get("profile", DocumentProfile())
    cfg = ctx.config
    warnings: List[str] = []
    errors: List[str] = []

    raw_blocks: List[Dict[str, Any]] = []
    raw_tables: List[Dict[str, Any]] = []
    raw_images: List[Dict[str, Any]] = []
    page_count = profile.page_count or 1

    try:
        if profile.file_type == "pdf":
            raw_blocks, raw_tables = _extract_pdf(path, profile, cfg, warnings)
        elif profile.file_type == "docx":
            raw_blocks, raw_tables = _extract_docx(path, profile, cfg, warnings)
        elif profile.file_type == "doc":
            raw_blocks, raw_tables = _extract_doc(path, profile, cfg, warnings)
        elif profile.file_type == "html":
            raw_blocks, raw_tables = _extract_html(path, profile, cfg, warnings)
        elif profile.file_type == "image":
            raw_blocks = _extract_image(path, profile, cfg, warnings)
        else:
            warnings.append(f"No extractor for file_type='{profile.file_type}'. Returning empty blocks.")
    except Exception as exc:
        errors.append(f"Extraction failed: {exc}")
        ctx.logger.error("Extraction error", stage="extraction", error=str(exc))

    ctx.put("raw_blocks", raw_blocks)
    ctx.put("raw_tables", raw_tables)
    ctx.put("raw_images", raw_images)

    return StageOutput(
        stage_name="extraction",
        status=STATUS_FAIL if errors else (STATUS_WARNING if warnings else STATUS_PASS),
        summary=f"Extracted {len(raw_blocks)} blocks, {len(raw_tables)} tables from {page_count} pages",
        warnings=warnings,
        errors=errors,
        output_summary={
            "raw_block_count": len(raw_blocks),
            "raw_table_count": len(raw_tables),
            "raw_image_count": len(raw_images),
        },
    )


# ---------------------------------------------------------------------------
# Extraction helpers
# ---------------------------------------------------------------------------


def _extract_pdf(
    path: Path, profile: DocumentProfile, cfg: PipelineConfig, warnings: List[str]
) -> Tuple[List[Dict], List[Dict]]:
    blocks: List[Dict] = []
    tables: List[Dict] = []
    try:
        import pdfminer.high_level as pdfminer
        from pdfminer.layout import LAParams

        for page_num, page_layout in enumerate(
            pdfminer.extract_pages(str(path), laparams=LAParams()), start=1
        ):
            for element_idx, element in enumerate(page_layout):
                if hasattr(element, "get_text"):
                    text = element.get_text().strip()
                    if text:
                        blocks.append({
                            "page_index": page_num,
                            "element_index": element_idx,
                            "raw_text": text,
                            "bbox": getattr(element, "bbox", None),
                            "source": "pdfminer",
                        })
    except ImportError:
        warnings.append("pdfminer not installed; PDF text extraction skipped.")
    except Exception as exc:
        warnings.append(f"PDF extraction error: {exc}")
    return blocks, tables


def _extract_docx(
    path: Path, profile: DocumentProfile, cfg: PipelineConfig, warnings: List[str]
) -> Tuple[List[Dict], List[Dict]]:
    blocks: List[Dict] = []
    tables: List[Dict] = []
    try:
        import docx

        doc = docx.Document(str(path))
        for para_idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                blocks.append({
                    "page_index": 1,
                    "element_index": para_idx,
                    "raw_text": para.text.strip(),
                    "style": para.style.name if para.style else "",
                    "source": "python-docx",
                })
        for tbl_idx, tbl in enumerate(doc.tables):
            rows = []
            header_row_index: Optional[int] = None
            for row_idx, row in enumerate(tbl.rows):
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
                # Detect header row: first row whose cells are all bold or use a heading style
                if header_row_index is None and row_idx == 0 and cells:
                    try:
                        first_cell = row.cells[0]
                        is_bold = any(
                            run.bold
                            for para in first_cell.paragraphs
                            for run in para.runs
                        )
                        if is_bold:
                            header_row_index = 0
                    except Exception:
                        pass
            tables.append({
                "page_index": 1,
                "table_index": tbl_idx,
                "rows": rows,
                "header_row_index": header_row_index,
                "source": "python-docx",
            })
    except ImportError:
        warnings.append("python-docx not installed; DOCX extraction skipped.")
    except Exception as exc:
        warnings.append(f"DOCX extraction error: {exc}")
    return blocks, tables


def _extract_doc(
    path: Path, profile: DocumentProfile, cfg: PipelineConfig, warnings: List[str]
) -> Tuple[List[Dict], List[Dict]]:
    """
    Extract text from a legacy .doc file.
    Uses _read_doc_text() which tries multiple backends in priority order.
    Tables: docx2txt does not preserve table topology, so we emit the
    raw text as blocks and add a warning about table loss.
    """
    blocks: List[Dict] = []
    tables: List[Dict] = []

    text = _read_doc_text(path, warnings)
    if not text:
        return blocks, tables

    # Normalize line endings: win32com returns \r as paragraph separators;
    # LibreOffice may return \r\n. Normalize to \n before splitting.
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Split on single newlines so each legal heading (Điều, Khoản, Chương…)
    # becomes its own block.
    raw_lines = [p.strip() for p in text.split("\n") if p.strip()]

    # Deduplicate short repeated lines (headers/footers repeated on each page).
    seen_short: Dict[str, int] = {}
    for idx, para in enumerate(raw_lines):
        normalized = re.sub(r"\s+", " ", para.lower())
        if len(normalized) < 80:
            count = seen_short.get(normalized, 0) + 1
            seen_short[normalized] = count
            if count > 1:
                continue  # skip repeated short header/footer
        blocks.append({
            "page_index": 1,
            "element_index": idx,
            "raw_text": para,
            "source": "doc_text_extraction",
        })

    warnings.append(
        f"'.doc' format: table topology cannot be reliably preserved with text extraction. "
        "Convert to .docx for full table support."
    )
    return blocks, tables


def _extract_html(
    path: Path, profile: DocumentProfile, cfg: PipelineConfig, warnings: List[str]
) -> Tuple[List[Dict], List[Dict]]:
    blocks: List[Dict] = []
    tables: List[Dict] = []
    try:
        from bs4 import BeautifulSoup

        html_text = path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(html_text, "html.parser")

        # Extract paragraphs and headings
        for el_idx, tag in enumerate(soup.find_all(["p", "h1", "h2", "h3", "h4", "h5", "li"])):
            text = tag.get_text(strip=True)
            if text:
                blocks.append({
                    "page_index": 1,
                    "element_index": el_idx,
                    "raw_text": text,
                    "tag": tag.name,
                    "source": "beautifulsoup",
                })

        # Extract tables
        for tbl_idx, tbl in enumerate(soup.find_all("table")):
            rows = []
            for row in tbl.find_all("tr"):
                rows.append([cell.get_text(strip=True) for cell in row.find_all(["td", "th"])])
            tables.append({
                "page_index": 1,
                "table_index": tbl_idx,
                "rows": rows,
                "source": "beautifulsoup",
            })
    except ImportError:
        # Fallback: plain text extraction
        warnings.append("beautifulsoup4 not installed; falling back to plain HTML text extraction.")
        text = path.read_text(encoding="utf-8", errors="replace")
        # Strip tags crudely
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        if text:
            blocks.append({
                "page_index": 1,
                "element_index": 0,
                "raw_text": text,
                "source": "regex_strip",
            })
    except Exception as exc:
        warnings.append(f"HTML extraction error: {exc}")
    return blocks, tables


def _extract_image(
    path: Path, profile: DocumentProfile, cfg: PipelineConfig, warnings: List[str]
) -> List[Dict]:
    blocks: List[Dict] = []
    if not cfg.enable_ocr:
        warnings.append("OCR disabled; image text extraction skipped.")
        return blocks
    try:
        import pytesseract
        from PIL import Image as PILImage

        img = PILImage.open(str(path))
        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        words = [
            w for w, conf in zip(data["text"], data["conf"])
            if w.strip() and int(conf) > 0
        ]
        confidences = [int(c) for c in data["conf"] if int(c) > 0]
        avg_conf = sum(confidences) / len(confidences) if confidences else 0
        text = " ".join(words)
        if text:
            blocks.append({
                "page_index": 1,
                "element_index": 0,
                "raw_text": text,
                "ocr_confidence": avg_conf / 100.0,
                "source": "tesseract",
            })
    except ImportError:
        warnings.append("pytesseract or Pillow not installed; OCR skipped for image input.")
    except Exception as exc:
        warnings.append(f"OCR error: {exc}")
    return blocks


# ---------------------------------------------------------------------------
# Shared .doc text reader (used by both profiler and extractor)
# ---------------------------------------------------------------------------


def _read_doc_text(path: Path, warnings: List[str]) -> str:
    """
    Extract plain text from a legacy .doc file.

    Priority order (first success wins):
    1. win32com — Word COM automation (Windows + MS Word; most reliable for OLE .doc)
    2. LibreOffice — headless CLI convert to text (cross-platform)
    3. docx2txt   — pure Python; works on XML-based .doc saved by newer Word; fails on OLE
    4. antiword   — Linux/macOS CLI
    5. python-docx — last resort; only works if file is secretly OOXML

    docx2txt uses ZIP/XML parsing and silently returns empty on true OLE .doc files,
    so win32com and LibreOffice are tried first on systems where they are available.
    """
    abs_path = str(path.resolve())

    # Attempt 1: win32com Word COM automation (Windows only, requires pywin32 + MS Word)
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(abs_path, ReadOnly=True)
        text = doc.Content.Text
        doc.Close(False)
        word.Quit()
        pythoncom.CoUninitialize()
        if text and text.strip():
            return text
    except ImportError:
        pass  # pywin32 not installed
    except Exception:
        pass  # COM error or Word not available

    # Attempt 2: LibreOffice headless CLI (soffice) — converts to .txt in a temp dir
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            result = subprocess.run(
                ["soffice", "--headless", "--convert-to", "txt:Text", "--outdir", tmp_dir, abs_path],
                capture_output=True, timeout=60,
            )
            if result.returncode == 0:
                txt_file = Path(tmp_dir) / (path.stem + ".txt")
                if txt_file.exists():
                    text = txt_file.read_text(encoding="utf-8", errors="replace")
                    if text.strip():
                        return text
    except FileNotFoundError:
        pass  # soffice not installed
    except Exception:
        pass

    # Attempt 3: docx2txt (pure Python; works on XML-based .doc, silently empty on OLE)
    try:
        import docx2txt
        text = docx2txt.process(abs_path)
        if text and text.strip():
            return text
    except ImportError:
        pass
    except Exception:
        pass

    # Attempt 4: antiword CLI (Linux/macOS)
    try:
        result = subprocess.run(
            ["antiword", abs_path],
            capture_output=True, text=True, timeout=30,
            encoding="utf-8", errors="replace",
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except FileNotFoundError:
        pass
    except Exception:
        pass

    # Attempt 5: python-docx (only succeeds if .doc is secretly OOXML)
    try:
        import docx
        doc = docx.Document(abs_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if text.strip():
            return text
    except Exception:
        pass

    warnings.append(
        f"Could not extract text from '{path.name}'. "
        "For .doc support on Windows: install pywin32 (`pip install pywin32`) with MS Word, "
        "or install LibreOffice and ensure `soffice` is on PATH. "
        "Alternatively convert the file to .docx before processing."
    )
    return ""
